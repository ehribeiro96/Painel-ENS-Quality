
# LEGADO: app Flask mantido apenas como referencia. Use run.py / fastapi_app.py.
"""
Gerador de Assinaturas ENS - Flask + SQLite
- Aviso padrao bilingue (PT/EN)
- Layout fiel da ENS (logo, cores, espacamento, icones sociais)
- Gestao de colaboradores (inclusao, edicao, remocao, ativacao)
- Exportacao individual (.html e .htm) e lote em ZIP
- Compatibilidade garantida para Windows 11 + Python 3.13 (debug=False, use_reloader=False)
"""

import base64
import datetime
import hmac
import io
import mimetypes
import os
import re
import secrets
import smtplib
import sqlite3
import sys
import unicodedata
import zipfile
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from html import escape, unescape
from pathlib import Path

from flask import Blueprint, Flask, abort, jsonify, make_response, redirect, render_template, request, send_file, send_from_directory, session, url_for
from markupsafe import Markup
from werkzeug.security import check_password_hash, generate_password_hash

from src.config.settings import (
    data_dir,
    default_db_path,
    env_flag,
    load_env_files,
    preview_cache_dir,
    static_dir,
    templates_dir,
)
from .assets_base64 import ASSETS as BASE64_ASSETS
from .signature_model_spec import (
    AVISO_TEXTO_EN,
    AVISO_TEXTO_PT,
    CANONICAL_HTML_REQUIRED_MARKERS,
    SIGNATURE_MODEL_SPEC,
    map_pt_to_html_font_size,
)

# Bibliotecas de terceiros usadas na integracao com Excel
import pandas as pd

# Imports relacionados ao OAuth2 / Microsoft Graph
import requests
try:
    # msal pode nao estar instalado em todos os ambientes ate ser definido no requirements.txt
    from msal import ConfidentialClientApplication
except ImportError:
    ConfidentialClientApplication = None  # type: ignore

try:
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE  # type: ignore
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.text.run import Run  # type: ignore
    from docx.opc.constants import RELATIONSHIP_TYPE  # type: ignore
except ImportError:
    Document = None  # type: ignore
    Inches = None  # type: ignore
    Pt = None  # type: ignore
    RGBColor = None  # type: ignore
    WD_ROW_HEIGHT_RULE = None  # type: ignore
    Run = None  # type: ignore
    RELATIONSHIP_TYPE = None  # type: ignore
    WD_TABLE_ALIGNMENT = None  # type: ignore
    WD_ALIGN_VERTICAL = None  # type: ignore
    WD_PARAGRAPH_ALIGNMENT = None  # type: ignore
    OxmlElement = None  # type: ignore
    qn = None  # type: ignore


def _read_secret_file(path):
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return f.read().strip()
    except Exception:
        pass
    return None


# Load secret values from secret files if provided (fallback to env)
ENS_SECRET_KEY = os.environ.get("ENS_SECRET_KEY") or _read_secret_file(os.environ.get("ENS_SECRET_KEY_FILE"))
ENS_MS_CLIENT_ID = os.environ.get("ENS_MS_CLIENT_ID") or _read_secret_file(os.environ.get("ENS_MS_CLIENT_ID_FILE"))
ENS_MS_CLIENT_SECRET = os.environ.get("ENS_MS_CLIENT_SECRET") or _read_secret_file(os.environ.get("ENS_MS_CLIENT_SECRET_FILE"))


def _ensure_db_env() -> None:
    """
    Garante que ENS_DB_PATH esteja definido apontando para data/ens.db quando
    nao houver configuracao explicita.
    """
    if os.getenv("ENS_DB_PATH"):
        return
    db_path = default_db_path()
    db_path.parent.mkdir(parents=True, exist_ok=True)
    os.environ["ENS_DB_PATH"] = str(db_path)


load_env_files()
_ensure_db_env()

_CSRF_SESSION_KEY = "_csrf_token"


def _get_or_create_csrf_token() -> str:
    """
    Retorna o token CSRF da sessão, gerando um novo quando inexistente.
    """
    token = session.get(_CSRF_SESSION_KEY)
    if not token:
        token = secrets.token_urlsafe(32)
        session[_CSRF_SESSION_KEY] = token
    return token


def _render_csrf_field() -> Markup:
    """
    Produz o campo hidden com o token CSRF atual.
    """
    token = _get_or_create_csrf_token()
    return Markup(f'<input type="hidden" name="csrf_token" value="{token}">')


def _is_secure_request() -> bool:
    """
    Detecta se a requisição atual é HTTPS diretamente ou por trás de proxy.
    """
    if request.is_secure:
        return True
    forwarded_proto = request.headers.get("X-Forwarded-Proto", "")
    if forwarded_proto:
        proto = forwarded_proto.split(",")[0].strip()
        if proto.lower() == "https":
            return True
    return False


DEFAULT_CSP = (
    "default-src 'self'; "
    "base-uri 'self'; "
    "form-action 'self'; "
    "img-src 'self' data: https:; "
    "font-src 'self' data: https:; "
    "style-src 'self' 'unsafe-inline' https:; "
    "script-src 'self' 'unsafe-inline'; "
    "connect-src 'self'; "
    "frame-ancestors 'none'"
)


def _apply_security_layers(aplicativo: Flask) -> None:
    """
    Aplica configuracoes de cookies, integracao CSRF e headers de seguranca.
    """
    same_site = os.getenv("ENS_SESSION_SAMESITE", "Lax")
    aplicativo.config.setdefault("SESSION_COOKIE_SAMESITE", same_site)
    aplicativo.config.setdefault("SESSION_COOKIE_HTTPONLY", env_flag("ENS_SESSION_HTTPONLY", True))
    aplicativo.config.setdefault("SESSION_COOKIE_SECURE", env_flag("ENS_SESSION_SECURE", True))
    aplicativo.config.setdefault("REMEMBER_COOKIE_HTTPONLY", True)
    aplicativo.config.setdefault("REMEMBER_COOKIE_SECURE", True)
    try:
        lifetime_hours = int(os.getenv("ENS_SESSION_HOURS", "8"))
    except ValueError:
        lifetime_hours = 8
    aplicativo.config.setdefault("PERMANENT_SESSION_LIFETIME", datetime.timedelta(hours=max(lifetime_hours, 1)))

    csp_policy = os.getenv("ENS_CONTENT_SECURITY_POLICY", DEFAULT_CSP)
    hsts_policy = os.getenv("ENS_HSTS_HEADER", "max-age=31536000; includeSubDomains")
    frame_options = os.getenv("ENS_FRAME_OPTIONS", "DENY")

    @aplicativo.context_processor
    def _inject_csrf_utilities():
        return {
            "csrf_token": _get_or_create_csrf_token,
            "csrf_field": _render_csrf_field,
        }

    @aplicativo.before_request
    def _enforce_csrf():
        if request.method in ("GET", "HEAD", "OPTIONS", "TRACE"):
            return
        token = session.get(_CSRF_SESSION_KEY)
        if not token:
            abort(400, description="Sess\u00e3o expirada. Recarregue a p\u00e1gina.")

        candidato = request.form.get("csrf_token")
        if not candidato:
            for header_name in ("X-CSRFToken", "X-CSRF-Token"):
                candidato = request.headers.get(header_name)
                if candidato:
                    break
        if not candidato and request.is_json:
            payload = request.get_json(silent=True) or {}
            candidato = payload.get("csrf_token")

        if not candidato or not hmac.compare_digest(str(token), str(candidato)):
            abort(400, description="Token CSRF inv\u00e1lido.")

    @aplicativo.after_request
    def _inject_security_headers(response):
        response.headers.setdefault("Content-Security-Policy", csp_policy)
        response.headers.setdefault("X-Frame-Options", frame_options)
        response.headers.setdefault("X-Content-Type-Options", "nosniff")
        response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
        response.headers.setdefault("Permissions-Policy", "geolocation=(), microphone=(), camera=()")
        if _is_secure_request():
            response.headers.setdefault("Strict-Transport-Security", hsts_policy)
        return response


TITULO_APLICACAO = "Gerador de Assinaturas ENS"
# ==============================
# Configuração do Banco de Dados
# ==============================
# A variável de ambiente `ENS_DB_PATH` permite apontar o arquivo `ens.db`
# para outro local (ex.: volume Docker, pasta compartilhada, etc.).
# Se ela não estiver definida, o projeto usa por padrão `ens.db` na pasta
# do próprio `app.py`.
if os.environ.get("ENS_DB_PATH"):
    CAMINHO_BANCO_DADOS = Path(os.environ.get("ENS_DB_PATH"))
else:
    CAMINHO_BANCO_DADOS = default_db_path()
IMAGEM_HEROI_PUBLICA = os.environ.get("ENS_PUBLIC_HERO", "")
HERO_PUBLICO_ATIVO = str(os.environ.get("ENS_PUBLIC_HERO_ENABLED", "off")).lower() in ("1", "true", "yes", "on")
ENDERECO_RIO_BASE = "Rua Senador Dantas, 74 | Centro - Rio de Janeiro - RJ - 20031-205"
ENDERECO_RIO = os.environ.get("ENS_ADDR_RIO", ENDERECO_RIO_BASE)
ENDERECO_SP = os.environ.get("ENS_ADDR_SP", "Rua Libero Badaro, 293 - 27o andar - Cj. 27D | Centro - Sao Paulo - SP - 01009-907")
ENDERECO_HOME = os.environ.get("ENS_ADDR_HOME", "Home Office")
_home_keywords_env = os.environ.get("ENS_HOME_KEYWORDS")
HOME_OFFICE_KEYWORDS = tuple(
    s.strip().lower()
    for s in (_home_keywords_env.split(",") if _home_keywords_env else ["home office", "home-office", "home", "remoto", "remota", "remote"])
    if s.strip()
)
ASSINATURA_NOME_PADRAO = "Assinatura_ENS"
HOME_OFFICE_STYLE = (
    "margin:0;"
    "line-height:115%;"
    "mso-line-height-rule:exactly;"
    "font-family:Verdana,Arial,sans-serif;"
    "font-size:10pt;"
    "font-style:italic;"
    "color:#009DB7;"
)

URL_LOGO = "https://funenseg-my.sharepoint.com/personal/estevao_quality_ens_edu_br/Documents/Static/Logo.png"
# URLs de icones padrao (usadas quando os arquivos locais nao estao presentes)
URL_ICONE_FACEBOOK_PADRAO  = "https://funenseg-my.sharepoint.com/personal/estevao_quality_ens_edu_br/Documents/Static/facebook.png"
URL_ICONE_INSTAGRAM_PADRAO  = "https://funenseg-my.sharepoint.com/personal/estevao_quality_ens_edu_br/Documents/Static/instagram.png"
URL_ICONE_LINKEDIN_PADRAO  = "https://funenseg-my.sharepoint.com/personal/estevao_quality_ens_edu_br/Documents/Static/linkedin.png"
URL_ICONE_TIKTOK_PADRAO  = "https://funenseg-my.sharepoint.com/personal/estevao_quality_ens_edu_br/Documents/Static/tiktok.png"
URL_ICONE_YOUTUBE_PADRAO  = "https://funenseg-my.sharepoint.com/personal/estevao_quality_ens_edu_br/Documents/Static/youtube.png"
ICONE_LINKS_SOCIAIS = {
    "facebook": "https://www.facebook.com/EscolaDeNegociosESeguros",
    "instagram": "https://www.instagram.com/oficial.ens/",
    "linkedin": "https://www.linkedin.com/school/5402791",
    "tiktok": "https://www.tiktok.com/@oficial.ens",
    "youtube": "https://www.youtube.com/channel/UCWKYHpO2GdJ7nfxuNx3XctQ",
}

LEGACY_ENVIRONMENT = (os.environ.get("ENS_ENV") or os.environ.get("ENVIRONMENT") or "local").strip().lower()
USUARIO_ADMINISTRADOR = os.environ.get("ENS_ADMIN_USER", "admin")
SENHA_ADMINISTRADOR = os.environ.get("ENS_ADMIN_PASS")
if LEGACY_ENVIRONMENT == "production" and not SENHA_ADMINISTRADOR:
    raise RuntimeError("ENS_ADMIN_PASS must be defined for legacy admin in production")
if ENS_SECRET_KEY:
    CHAVE_SECRETA = ENS_SECRET_KEY
elif LEGACY_ENVIRONMENT == "production":
    raise RuntimeError("ENS_SECRET_KEY must be defined for legacy signatures in production")
else:
    CHAVE_SECRETA = secrets.token_urlsafe(32)
INCORPORAR_BASE64 = str(os.environ.get("ENS_EMBED_BASE64", "1")).lower() in ("1","true","yes","on")

# ------------------------------------------------------------------

# ===============================
# Envio de E-mails (modo híbrido)
# ===============================
# A aplicação pode enviar a assinatura por dois caminhos:
#   1) Microsoft Graph (OAuth2) -> uso padrão em produção
#   2) SMTP clássico            -> útil em homologação/apresentação
#
# A variável `ENS_SMTP_MODE` controla o comportamento:
#   - "ON"  -> força o uso de SMTP (ignora Graph)
#   - "OFF" -> usa o fluxo padrão via Microsoft Graph
EMAIL_SMTP_MODE = (os.environ.get("ENS_SMTP_MODE") or "OFF").strip().upper()
USE_SMTP_EMAIL = EMAIL_SMTP_MODE == "ON"

# Credenciais SMTP de homologação
# IMPORTANTE:
#   - Estes valores são apenas para ambiente de teste/apresentação.
#   - Em produção, use sempre variáveis de ambiente seguras.
# Variáveis suportadas:
#   ENS_SMTP_USER, ENS_SMTP_PASSWORD, ENS_SMTP_HOST, ENS_SMTP_PORT
SMTP_HOST = os.environ.get("ENS_SMTP_HOST", "smtp.office365.com")
SMTP_PORT = int(os.environ.get("ENS_SMTP_PORT", "587"))
SMTP_USER = os.environ.get("ENS_SMTP_USER")
SMTP_PASSWORD = os.environ.get("ENS_SMTP_PASSWORD")
if USE_SMTP_EMAIL and (not SMTP_USER or not SMTP_PASSWORD):
    raise RuntimeError("ENS_SMTP_USER and ENS_SMTP_PASSWORD must be defined when ENS_SMTP_MODE=ON")

# Configuracao Microsoft OAuth (Graph API)
# As variaveis abaixo devem estar definidas para que o OAuth2 funcione.
# Se alguma estiver ausente, login e envio via Outlook ficarao indisponiveis.
MS_ID_CLIENTE = ENS_MS_CLIENT_ID
MS_SEGREDO_CLIENTE = ENS_MS_CLIENT_SECRET
MS_ID_LOCATARIO = os.environ.get("ENS_MS_TENANT_ID", "common")
# Opcional: substitui o caminho de redirecionamento; por padrao utilizamos '/auth/callback'
MS_CAMINHO_REDIRECIONAMENTO = os.environ.get("ENS_MS_REDIRECT_PATH", "/auth/callback")
MS_ESCOPOS = os.environ.get("ENS_MS_SCOPES", "Mail.Send offline_access")
MS_IDS_GRUPOS_OBRIGATORIOS = [s.strip() for s in (os.environ.get("ENS_MS_REQUIRED_GROUP_IDS") or "").split(",") if s.strip()]

DIRETORIO_ESTATICO = static_dir().resolve()
TEMPLATE_ROOT = templates_dir().resolve()
DIRETORIO_ICONES = DIRETORIO_ESTATICO / "icons"
DIRETORIO_DATA = data_dir().resolve()
ARQUIVO_MODELO_WORD = DIRETORIO_ESTATICO / "ASSINATURAS DE E-MAIL (ENS_LOGO_AZUL_LGPD_semTWITTER)_v21.23.docx"
DIRETORIO_PREVIEWS = preview_cache_dir().resolve()
GUIA_ILUSTRADO_DOC = DIRETORIO_ESTATICO / "Guia_Assinatura_ENS_Ilustrado_v1.docx"
IDENTIFICADOR_ENDERECO_RIO = "Rua Senador Dantas, 74"
PADROES_ANDAR_RIO = [
    (
        re.compile(r"(?P<andar>\d+)\s*(?:\u00ba|o)?\s*andar", re.IGNORECASE),
        lambda m: f"{int(m.group('andar'))}\u00ba andar",
    ),
    (re.compile(r"t[e\u00e9]rre?o", re.IGNORECASE), lambda _m: "T\u00e9rreo"),
]

_GUIA_IMAGENS_CACHE: dict[str, bytes] = {}
_GUIA_ILUSTRADO_CACHE: tuple[str, list[tuple[str, str, bytes]]] | None = None

def criar_aplicativo(tipo: str = "ambos") -> Flask:
    tipo_normalizado = (tipo or "ambos").strip().lower()
    if tipo_normalizado not in ("admin", "publico", "ambos"):
        tipo_normalizado = "ambos"

    aplicativo = Flask(
        __name__,
        template_folder=str(TEMPLATE_ROOT),
        static_folder=str(DIRETORIO_ESTATICO),
    )
    aplicativo.config["SECRET_KEY"] = CHAVE_SECRETA
    aplicativo.config["TEMPLATES_AUTO_RELOAD"] = True
    aplicativo.config["ENS_APP_MODE"] = tipo_normalizado
    # Usa cookies de sessao distintos conforme o tipo de aplicativo
    if tipo_normalizado == "admin":
        aplicativo.config["SESSION_COOKIE_NAME"] = "ens_admin_session"
    elif tipo_normalizado == "publico":
        aplicativo.config["SESSION_COOKIE_NAME"] = "ens_public_session"

    _apply_security_layers(aplicativo)

    # Define sinais globais para a interface
    aplicativo_publico = tipo_normalizado in ("publico", "ambos")
    aplicativo_admin = tipo_normalizado in ("admin", "ambos")
    aplicativo.jinja_env.globals["APLICATIVO_PUBLICO"] = aplicativo_publico
    aplicativo.jinja_env.globals["APLICATIVO_ADMIN"] = aplicativo_admin

    # Registra os blueprints conforme o tipo
    if aplicativo_publico:
        aplicativo.register_blueprint(bp_publico)
    if aplicativo_admin:
        admin_kwargs = {} if tipo_normalizado == "admin" else {"url_prefix": "/admin"}
        aplicativo.register_blueprint(bp_admin, **admin_kwargs)

    # Garante que o banco esteja inicializado
    with aplicativo.app_context():
        inicializar_banco()

    # Se houver uma build de SPA (React) em `static/index.html`, sirva-a
    @aplicativo.route('/', defaults={'path': ''})
    @aplicativo.route('/<path:path>')
    def _serve_spa(path: str):
        try:
            index_path = DIRETORIO_ESTATICO / 'index.html'
            requested = DIRETORIO_ESTATICO / path
            # Serve arquivos estáticos existentes diretamente
            if path and requested.exists():
                return send_from_directory(str(DIRETORIO_ESTATICO), path)
            # Se existir index.html no diretório static, devolve-o para rotas do SPA
            if index_path.exists():
                return send_from_directory(str(DIRETORIO_ESTATICO), 'index.html')
        except Exception:
            pass
        # Não interferir em outras rotas/erros se não houver build
        abort(404)

    return aplicativo

def _normalizar_token(value: str | None) -> str:
    """
    Normaliza texto removendo acentos e convertendo para minusculas.
    Retorna string vazia quando nao ha valor informado.
    """
    if not value:
        return ""
    text = str(value).strip()
    try:
        text = "".join(c for c in unicodedata.normalize("NFKD", text) if not unicodedata.combining(c))
    except Exception:
        text = text.lower()
    return text.lower()

NOMES_MULTIPLAS_ASSINATURAS = {"paola casado"}


def _permite_assinatura_dupla(nome: str | None) -> bool:
    return _normalizar_token(nome) in NOMES_MULTIPLAS_ASSINATURAS


def eh_colaborador_atn(registro: dict | None) -> bool:
    if not registro:
        return False
    matricula = (registro.get("matricula") or "").strip().upper()
    if matricula.startswith("ATN"):
        return True
    departamento = (registro.get("department") or registro.get("diretoria") or "").strip().upper()
    return departamento == "ATN"


def _selecionar_registro_existente(
    rows: list[sqlite3.Row],
    email: str | None,
    nome: str | None,
) -> dict | None:
    if not rows:
        return None
    registros = [dict(row) for row in rows]
    email_norm = (email or "").strip().lower()
    if email_norm:
        for row in registros:
            if (row.get("email") or "").strip().lower() == email_norm:
                return row
    if _permite_assinatura_dupla(nome):
        for row in registros:
            if not _permite_assinatura_dupla(row.get("name")):
                return row
        return None
    return registros[0]

def classificar_area_por_matricula(matricula: str | None) -> str:
    """
    Classifica colaboradores em area ENS ou ATN a partir do prefixo da matricula.
    Matriculas vazias sao agrupadas como "Sem Matr\u00edcula" para facilitar a visualizacao.
    """
    token = (matricula or "").strip().upper()
    if not token:
        return "Sem Matr\u00edcula"
    if token.startswith("ATN"):
        return "ATN"
    if token.startswith("ENS"):
        return "ENS"
    # Assume ENS para matriculas numericas ou prefixos desconhecidos.
    return "ENS"

def extrair_rotulo_andar_rio(*texts: str | None) -> str | None:
    for text in texts:
        if not text:
            continue
        for pattern, label in PADROES_ANDAR_RIO:
            match = pattern.search(text)
            if not match:
                continue
            if callable(label):
                try:
                    return label(match)
                except Exception:
                    continue
            return label
    return None

def _inserir_andar_no_endereco_rio(base: str, floor: str) -> str:
    """
    Insere o rotulo do andar logo apos o numero (antes do bairro) replicando a diagramacao do Word.
    Caso o andar ja esteja presente, retorna o endereco original.
    """
    base_limpo = (base or "").strip()
    floor_limpo = (floor or "").strip()
    if not base_limpo or not floor_limpo:
        return base_limpo
    if floor_limpo.lower() in base_limpo.lower():
        return base_limpo
    marcador = " | "
    if marcador in base_limpo:
        primeira, resto = base_limpo.split(marcador, 1)
        primeira = primeira.rstrip()
        resto = resto.lstrip()
        return f"{primeira} - {floor_limpo}{marcador}{resto}"
    return f"{base_limpo} - {floor_limpo}"


def normalizar_endereco_rio(address: str | None, *hints: str | None) -> str | None:
    """
    Padroniza enderecos da Matriz do Rio utilizando uma base fixa e o andar identificado.
    Caso nao seja um endereco da Rua Senador Dantas, o valor informado e mantido sem alteracoes.
    """
    text = (address or "").strip()
    if text:
        # Mantem exatamente o endereco informado no banco (apenas remove espacos excedentes)
        return text
    valid_hints = [hint for hint in hints if hint]
    if not valid_hints:
        return None
    identifier = IDENTIFICADOR_ENDERECO_RIO.lower()
    identifier_present = any(identifier in hint.lower() for hint in valid_hints)
    if not identifier_present:
        return None
    floor = extrair_rotulo_andar_rio(*valid_hints)
    base = ENDERECO_RIO_BASE
    if floor:
        return _inserir_andar_no_endereco_rio(base, floor)
    return base


def inferir_tipo_endereco(registro: dict) -> str:
    """
    Determina o tipo de endereco (rio, sp, home) a partir dos dados armazenados.
    Utilizado para preencher o formulario publico com o valor coerente do banco.
    """
    uf = (registro.get("uf") or "").strip().upper()
    local_desc = (registro.get("local_descricao") or "").strip().lower()
    endereco = (registro.get("endereco") or "").strip().lower()
    texto = " ".join(filter(None, (local_desc, endereco)))
    if any(keyword in texto for keyword in HOME_OFFICE_KEYWORDS or ("home", "remoto")):
        return "home"
    if uf == "SP" or "sao paulo" in texto:
        return "sp"
    return "rio"


def resolver_home_office(dados: dict, tipo_inferido: str | None = None, *fontes_adicionais: str | None) -> tuple[bool, str]:
    """
    Determina se o colaborador atua em Home Office e retorna o rotulo configurado.
    Considera address_type, inferencia automatica e palavras-chave configuraveis.
    """
    keywords = HOME_OFFICE_KEYWORDS or ("home", "home office", "remoto")
    candidatos: list[str] = []
    for chave in ("local_descricao", "endereco", "address"):
        valor = dados.get(chave)
        if valor:
            candidatos.append(str(valor))
    for extra in fontes_adicionais:
        if extra:
            candidatos.append(str(extra))

    label = ""
    for texto in candidatos:
        texto_normalizado = texto.strip()
        if not texto_normalizado:
            continue
        texto_lower = texto_normalizado.lower()
        if any(keyword in texto_lower for keyword in keywords):
            label = texto_normalizado
            break

    address_type = (dados.get("address_type") or "").strip().lower()
    tipo_normalizado = (tipo_inferido or "").strip().lower()
    if not tipo_normalizado:
        tipo_normalizado = inferir_tipo_endereco(dados).strip().lower()
    is_home = address_type in ("home", "homeoffice", "home-office") or tipo_normalizado == "home"
    if label:
        is_home = True
    if is_home and not label:
        label = (ENDERECO_HOME or "").strip() or "Home Office"
    return is_home, label


def _formatar_data_brasil(valor: str | None) -> str:
    if not valor:
        return ""
    texto = valor.strip()
    if not texto:
        return ""
    candidato = texto.replace("Z", "+00:00")
    dt_obj = None
    try:
        dt_obj = datetime.datetime.fromisoformat(candidato)
    except ValueError:
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
            try:
                dt_obj = datetime.datetime.strptime(texto, fmt)
                break
            except ValueError:
                continue
    if not dt_obj:
        return texto
    return dt_obj.strftime("%d/%m/%Y %H:%M")


def _preparar_colaborador_publico(registro: dict) -> tuple[dict[str, object], dict[str, object]]:
    dados = dict(registro)
    nome_assinatura = (dados.get("nome_exibicao") or dados.get("name") or "").strip()
    if not nome_assinatura:
        nome_assinatura = dados.get("name") or ""

    tipo_inferido = inferir_tipo_endereco(dados)
    eh_home, home_label = resolver_home_office(dados, tipo_inferido, dados.get("endereco"))
    endereco_normalizado = normalizar_endereco_rio(dados.get("endereco"), dados.get("local_descricao"))

    endereco_display = endereco_normalizado or (dados.get("endereco") or "")
    address_label = ""
    if eh_home:
        address_label = home_label or ""
        endereco_display = home_label or ENDERECO_HOME or endereco_display

    colaborador_publico = {
        "nome": dados.get("name") or "",
        "nome_exibicao": nome_assinatura,
        "matricula": dados.get("matricula") or "",
        "email": dados.get("email") or "",
        "departamento": dados.get("department") or "",
        "cargo": dados.get("role") or "",
        "diretoria": dados.get("diretoria") or "",
        "posicao_organograma": dados.get("posicao_organograma") or "",
        "manager": dados.get("manager") or "",
        "telefone": dados.get("phone") or "",
        "telefone_ad": dados.get("telefone_ad") or "",
        "local_descricao": dados.get("local_descricao") or "",
        "endereco": endereco_display or "",
        "address_label": address_label,
        "uf": dados.get("uf") or "",
        "updated_at": _formatar_data_brasil(dados.get("updated_at")),
        "home_office": eh_home,
    }

    dados_assinatura = {
        "name": nome_assinatura or dados.get("name") or "",
        "nome_exibicao": nome_assinatura or dados.get("nome_exibicao") or "",
        "department": dados.get("department") or "",
        "role": dados.get("role") or "",
        "phone": dados.get("phone") or "",
        "email": dados.get("email") or "",
        "local_descricao": dados.get("local_descricao") or "",
        "endereco": dados.get("endereco") or "",
        "address_type": dados.get("address_type") or tipo_inferido,
    }

    if eh_home:
        dados_assinatura["address_type"] = "home"
        if address_label:
            dados_assinatura["endereco"] = address_label
            dados_assinatura["address"] = address_label
    elif endereco_normalizado:
        dados_assinatura["endereco"] = endereco_normalizado

    return colaborador_publico, dados_assinatura


def _armazenar_colaborador_lookup(
    colaborador_publico: dict[str, object],
    dados_assinatura: dict[str, object],
    registro_id: int | None = None,
) -> str:
    token = secrets.token_urlsafe(24)
    session["colaborador_lookup"] = {
        "token": token,
        "colaborador": colaborador_publico,
        "dados_assinatura": dados_assinatura,
        "registro_id": registro_id,
    }
    session.modified = True
    return token


# Campos editaveis/gravaveis via painel administrativo.
CAMPOS_COLABORADOR_EDITAVEIS: tuple[str, ...] = (
    "matricula",
    "name",
    "department",
    "role",
    "phone",
    "email",
    "posicao_organograma",
    "first_name",
    "last_name",
    "nome_exibicao",
    "diretoria",
    "campo_assinatura",
    "manager",
    "telefone_ad",
    "local_descricao",
    "endereco",
    "uf",
)


def coletar_dados_colaborador(form_data) -> dict[str, str | None]:
    """
    Converte valores do formulario em um dicionario pronto para persistencia.
    Remove espacos excedentes e traduz strings vazias para None.
    """
    dados: dict[str, str | None] = {}
    for campo in CAMPOS_COLABORADOR_EDITAVEIS:
        valor = form_data.get(campo)
        if valor is None:
            dados[campo] = None
            continue
        texto = valor.strip()
        dados[campo] = texto if texto else None
    return dados


def ajustar_cargo_com_departamento(dados: dict[str, str | None]) -> None:
    """
    Garante que o cargo nao fique vazio ou apenas com ponto quando houver departamento preenchido.
    """
    if eh_colaborador_atn(dados):
        return
    departamento = dados.get("department")
    if not departamento:
        return
    cargo = dados.get("role")
    if cargo is None:
        dados["role"] = departamento
        return
    cargo_limpo = cargo.strip()
    if not cargo_limpo or cargo_limpo == ".":
        dados["role"] = departamento

# ----------------------------------------------------------------------
# Auxiliar de sincronizacao com Azure AD (Graph API)

def importar_usuarios_do_graph() -> int:
    """
    Consulta usuarios no Microsoft Graph utilizando token de aplicacao e sincroniza a tabela de colaboradores.
    Retorna o total de registros importados ou atualizados. Requer as variaveis ENS_MS_CLIENT_ID,
    ENS_MS_CLIENT_SECRET e ENS_MS_TENANT_ID configuradas. Caso nao seja possivel obter o token,
    a funcao retorna 0 sem efetuar alteracoes.
    """
    # Cria aplicacao de credenciais para obter token somente de aplicacao
    app_instance = construir_aplicativo_msal()
    if app_instance is None:
        return 0
    try:
        result = app_instance.acquire_token_for_client(
            scopes=["https://graph.microsoft.com/.default"],
        )
    except Exception:
        return 0
    access_token = result.get("access_token")
    if not access_token:
        return 0
    headers = {"Authorization": f"Bearer {access_token}"}
    json_headers = dict(headers)
    json_headers["Content-Type"] = "application/json"
    # URL base com propriedades selecionadas; employeeId corresponde a matricula quando disponivel.
    base_url = (
        "https://graph.microsoft.com/v1.0/users"
        "?$select=id,displayName,givenName,surname,jobTitle,department,mail,userPrincipalName,businessPhones,mobilePhone,employeeId"
    )
    next_url: str | None = base_url
    total_importado = 0
    # Usa uma unica conexao durante toda a importacao para reduzir overhead
    with obter_conexao() as conn:
        while next_url:
            try:
                resp = requests.get(next_url, headers=headers, timeout=30)
            except Exception:
                break
            if resp.status_code != 200:
                break
            data = resp.json()
            users = data.get("value") or []
            for user in users:
            # Extrai campos com valores de fallback
                display_name = user.get("displayName") or ""
                first_name = user.get("givenName") or ""
                last_name = user.get("surname") or ""
                department = user.get("department") or ""
                role = user.get("jobTitle") or ""
                email = user.get("mail") or user.get("userPrincipalName") or ""
            # businessPhones e uma lista; utiliza o primeiro valor quando existir
                phones = user.get("businessPhones") or []
                phone = phones[0] if isinstance(phones, list) and phones else ""
                mobile = user.get("mobilePhone") or ""
                phone_val = phone or mobile
                matricula = user.get("employeeId") or ""
                if MS_IDS_GRUPOS_OBRIGATORIOS:
                    azure_id = user.get("id")
                    if not azure_id:
                        continue
                    try:
                        membership = requests.post(
                            f"https://graph.microsoft.com/v1.0/users/{azure_id}/checkMemberGroups",
                            headers=json_headers,
                            json={"groupIds": MS_IDS_GRUPOS_OBRIGATORIOS},
                            timeout=30,
                        )
                    except Exception:
                        continue
                    if membership.status_code != 200:
                        continue
                    matched_groups = membership.json().get("value") or []
                    if not matched_groups:
                        continue
            # Define o nome de exibicao: prefere displayName do AD; senao monta a partir de nome e sobrenome
                name_val = display_name if display_name else (first_name + (" " + last_name if last_name else ""))
            # Ignora registros sem e-mail (Graph pode retornar contas de servico sem endereco)
                if not email:
                    continue
            # Busca registro existente por e-mail ou matricula para preservar status e senhas
                existing_rows = conn.execute(
                    "SELECT * FROM colaboradores WHERE email=? OR matricula=?",
                    (email, matricula if matricula else email),
                ).fetchall()
                existing = _selecionar_registro_existente(existing_rows, email, name_val)
                existing_id = existing.get("id") if existing else None
                status_val = (existing.get("status") if existing else None) or "on"
                eh_admin_val = (existing.get("eh_admin") if existing else None) or "no"
                password_hash_val = existing.get("password_hash") if existing else None
                must_change_val = existing.get("must_change") if existing else "no"
            # Mantem a matricula anterior caso o registro do AD nao informe esse dado
                if matricula:
                    matricula_val = matricula
                else:
                    matricula_val = existing.get("matricula") if existing else None
                updated_at = datetime.datetime.now().isoformat(timespec="seconds")
            # Prepara a instrucao de insercao/substituicao. Apenas campos vindos do AD sao atualizados;
            # campos adicionais (diretoria, campo_assinatura etc.) permanecem inalterados.
                conn.execute(
                    "INSERT OR REPLACE INTO colaboradores (id, matricula, name, department, role, phone, email, status, updated_at, eh_admin, password_hash, must_change, first_name, last_name) "
                    "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (
                        existing_id,
                        matricula_val,
                        name_val,
                        department,
                        role,
                        phone_val,
                        email,
                        status_val,
                        updated_at,
                        eh_admin_val,
                        password_hash_val,
                        must_change_val,
                        first_name,
                        last_name,
                    ),
                )
                total_importado += 1
        # Avanca para a proxima pagina quando houver
            next_url = data.get("@odata.nextLink")
    return total_importado


# ----- Funcoes auxiliares do banco -----
def padronizar_enderecos_matriz(conn: sqlite3.Connection) -> None:
    """
    Padroniza os enderecos armazenados da Matriz RJ para que todas as variacoes
    utilizem a mesma base e sufixo de andar. Evita duplicar registros semelhantes
    (3o, 4o andar, Terreo) e mantem as assinaturas consistentes.
    """
    try:
        rows = conn.execute(
            "SELECT id, endereco, local_descricao FROM colaboradores WHERE endereco LIKE ? OR (local_descricao IS NOT NULL AND local_descricao LIKE ?)",
            (f"%{IDENTIFICADOR_ENDERECO_RIO}%", f"%{IDENTIFICADOR_ENDERECO_RIO}%"),
        ).fetchall()
    except Exception:
        return
    for row in rows:
        current = row["endereco"]
        normalized = normalizar_endereco_rio(current, row["local_descricao"])
        if normalized and normalized != current:
            conn.execute("UPDATE colaboradores SET endereco=? WHERE id=?", (normalized, row["id"]))


def sincronizar_cargos_departamento(conn: sqlite3.Connection) -> None:
    """
    Atualiza cargos vazios ou preenchidos apenas com ponto para refletirem o departamento correspondente.
    """
    try:
        conn.execute(
            """
            UPDATE colaboradores
               SET role = TRIM(department)
             WHERE (role IS NULL OR TRIM(role) = '' OR TRIM(role) = '.')
               AND department IS NOT NULL
               AND TRIM(department) <> ''
            """
        )
    except Exception:
        pass

def obter_conexao():
    # Timeout maior aumenta resiliencia em escritas concorrentes; WAL melhora concorrencia entre processos
    conn = sqlite3.connect(CAMINHO_BANCO_DADOS, timeout=30)
    conn.row_factory = sqlite3.Row
    try:
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA synchronous=NORMAL")
    except Exception:
        pass
    return conn

def inicializar_banco():
    with obter_conexao() as conn:
        conn.execute(
            "CREATE TABLE IF NOT EXISTS colaboradores ("
            "id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "name TEXT,"
            "department TEXT,"
            "role TEXT,"
            "phone TEXT,"
            "email TEXT UNIQUE,"
            "status TEXT DEFAULT 'on',"
            "updated_at TEXT"
            ")"
        )
        # Migracao: garantir coluna 'eh_admin' (renomeando 'is_admin' quando existir)
        colunas = [linha["name"] for linha in conn.execute("PRAGMA table_info(colaboradores)").fetchall()]
        if "eh_admin" not in colunas:
            if "is_admin" in colunas:
                try:
                    conn.execute("ALTER TABLE colaboradores RENAME COLUMN is_admin TO eh_admin")
                except Exception:
                    pass
            else:
                try:
                    conn.execute("ALTER TABLE colaboradores ADD COLUMN eh_admin TEXT DEFAULT 'no'")
                except Exception:
                    pass
        # Migracao: adicionar password_hash e must_change para usuarios administradores
        try:
            conn.execute("ALTER TABLE colaboradores ADD COLUMN password_hash TEXT")
        except Exception:
            pass
        try:
            conn.execute("ALTER TABLE colaboradores ADD COLUMN must_change TEXT DEFAULT 'no'")
        except Exception:
            pass
        # Migracao: adicionar campo matricula (texto simples). O SQLite nao permite adicionar coluna UNIQUE via ALTER TABLE,
        # entao criamos uma coluna TEXT simples e posteriormente criamos um indice, se necessario.
        try:
            conn.execute("ALTER TABLE colaboradores ADD COLUMN matricula TEXT")
        except Exception:
            pass
        # Cria indice em matricula para acelerar buscas (nao unico para evitar erro caso haja duplicados)
        try:
            conn.execute("CREATE INDEX IF NOT EXISTS idx_colaboradores_matricula ON colaboradores(matricula)")
        except Exception:
            pass

        # ------------------------------------------------------------------
        # Colunas adicionais para refletir campos da planilha do RH.
        # Cada ALTER TABLE fica em try/except para que inicializar_banco possa rodar varias vezes
        # sem gerar excecao caso a coluna ja exista.
        additional_columns = [
            ("posicao_organograma", "TEXT"),
            ("first_name", "TEXT"),
            ("last_name", "TEXT"),
            ("nome_exibicao", "TEXT"),
            ("diretoria", "TEXT"),
            ("campo_assinatura", "TEXT"),
            ("manager", "TEXT"),
            ("telefone_ad", "TEXT"),
            ("local_descricao", "TEXT"),
            ("endereco", "TEXT"),
            ("uf", "TEXT"),
        ]
        for col_name, col_type in additional_columns:
            try:
                conn.execute(f"ALTER TABLE colaboradores ADD COLUMN {col_name} {col_type}")
            except Exception:
                # Coluna possivelmente existente; ignora erros
                pass

        padronizar_enderecos_matriz(conn)
        sincronizar_cargos_departamento(conn)

# ----- Resolucao de recursos (icones) -----
def _gerar_data_uri(path: Path) -> str | None:
    try:
        if not path.exists():
            return None
        data = path.read_bytes()
        # Identifica o tipo de conteudo; padrao png
        mime, _ = mimetypes.guess_type(str(path))
        if not mime:
            mime = "image/png"
        b64 = base64.b64encode(data).decode("ascii")
        return f"data:{mime};base64,{b64}"
    except Exception:
        return None

def _extrair_bytes_data_uri(payload: str | None) -> bytes | None:
    """
    Decodifica uma data URI (data:image/png;base64,...) e retorna os bytes puros.
    """
    if not payload or not payload.startswith("data:"):
        return None
    try:
        header, data = payload.split(",", 1)
    except ValueError:
        return None
    if ";base64" not in header.lower():
        return None
    try:
        return base64.b64decode(data)
    except Exception:
        return None

def _asset_fallback(name: str) -> str:
    """
    Fallback para recursos base64 embarcados no codigo.
    Retorna string vazia quando o recurso nao estiver mapeado.
    """
    return BASE64_ASSETS.get(name.lower(), "")

def _armazenar_preview_html(html: str) -> str:
    """
    Salva o HTML em disco e retorna um token curto para recupera-lo depois.
    Evita armazenar HTML volumoso no cookie de sessao.
    """
    DIRETORIO_PREVIEWS.mkdir(parents=True, exist_ok=True)
    token = secrets.token_hex(16)
    destino = DIRETORIO_PREVIEWS / f"{token}.html"
    try:
        destino.write_text(html, encoding="utf-8")
    except Exception:
        return ""
    return token

def _ler_preview_html(token: str | None) -> str | None:
    if not token:
        return None
    caminho = DIRETORIO_PREVIEWS / f"{token}.html"
    if not caminho.exists():
        return None
    try:
        return caminho.read_text(encoding="utf-8")
    except Exception:
        return None


def _bytes_para_data_uri(data: bytes, mime: str) -> str:
    b64 = base64.b64encode(data).decode("ascii")
    return f"data:{mime};base64,{b64}"


def _carregar_zip_entry_bytes(arquivo: Path, membro: str) -> bytes | None:
    try:
        with zipfile.ZipFile(arquivo) as zf:
            return zf.read(membro)
    except Exception:
        return None


def _imagem_guia_bytes(nome: str) -> bytes | None:
    cache = _GUIA_IMAGENS_CACHE
    if nome in cache:
        return cache[nome]
    if not GUIA_ILUSTRADO_DOC.exists():
        return None
    dados = _carregar_zip_entry_bytes(GUIA_ILUSTRADO_DOC, f"word/media/{nome}")
    if not dados:
        return None
    cache[nome] = dados
    return dados


def _imagem_guia_html(nome: str, legenda: str, anexos: list[tuple[str, str, bytes]]) -> str:
    dados = _imagem_guia_bytes(nome)
    if not dados:
        return ""
    cid = f"guia-{nome}"
    anexos.append((cid, "image/png", dados))
    return (
        "<div style='margin:16px 0;text-align:center;'>"
        f"<img src='cid:{cid}' alt='{legenda}' "
        "style='max-width:100%;border-radius:12px;box-shadow:0 8px 24px rgba(0,0,0,0.12);'>"
        "</div>"
    )


def _html_guia_assinatura_ilustrado() -> tuple[str, list[tuple[str, str, bytes]]] | None:
    global _GUIA_ILUSTRADO_CACHE
    if _GUIA_ILUSTRADO_CACHE is not None:
        html, anexos = _GUIA_ILUSTRADO_CACHE
        return html, [(cid, mime, data) for cid, mime, data in anexos]
    if not GUIA_ILUSTRADO_DOC.exists():
        _GUIA_ILUSTRADO_CACHE = ("", [])
        return None
    try:
        anexos: list[tuple[str, str, bytes]] = []
        partes: list[str] = []
        partes.append("<div style='font-family:\"Segoe UI\",Arial,sans-serif;color:#0b2630;line-height:1.55;font-size:14px;'>")
        partes.append("<h2 style='color:#00687a;margin-top:0;'>Guia ilustrado - Como instalar a Assinatura ENS no Outlook</h2>")
        partes.append(
            "<p>Este guia reproduz o documento oficial <em>Guia_Assinatura_ENS_Ilustrado_v1</em>, com os mesmos passos e ilustrações para Outlook clássico e Outlook Web.</p>"
        )
        partes.append("<h3 style='color:#009db7;margin-bottom:8px;'>Parte 1 - Outlook clássico para Windows</h3>")
        partes.append("<h4 style='margin-bottom:4px;'>Passo 1 - Acesse Arquivo &gt; Opções &gt; Email &gt; Assinaturas</h4>")
        partes.append(
            "<ol style='padding-left:20px;margin-top:4px;'>"
            "<li>Abra o Outlook instalado no computador ENS.</li>"
            "<li>Clique em <strong>Arquivo</strong> no canto superior esquerdo.</li>"
            "<li>Em seguida, selecione <strong>Opções</strong>.</li>"
            "<li>Na janela <strong>Opções do Outlook</strong>, clique na aba <strong>Email</strong>.</li>"
            "<li>Clique no botão <strong>Assinaturas...</strong> para abrir a tela de assinaturas.</li>"
            "</ol>"
        )
        partes.append(_imagem_guia_html("image1.png", "Arquivo > Opções > Email > Assinaturas", anexos))

        partes.append("<h4 style='margin-bottom:4px;'>Passo 2 - Criar a assinatura “Assinatura ENS”</h4>")
        partes.append(
            "<ol style='padding-left:20px;margin-top:4px;'>"
            "<li>Na janela <strong>Assinaturas e papel de carta</strong>, clique em <strong>Novo</strong>.</li>"
            "<li>Digite um nome como <strong>Assinatura ENS</strong>.</li>"
            "<li>Cole (Ctrl+V) a assinatura copiada do Portal de Assinaturas ENS no quadro de edição.</li>"
            "<li>Confirme se o texto e a formatação estão iguais ao modelo ENS.</li>"
            "</ol>"
        )
        partes.append(_imagem_guia_html("image2.png", "Janela de configuração da assinatura ENS", anexos))

        partes.append("<h4 style='margin-bottom:4px;'>Passo 3 - Conferir e salvar a assinatura</h4>")
        partes.append(
            "<ol style='padding-left:20px;margin-top:4px;'>"
            "<li>Verifique se a assinatura mostra corretamente nome, cargo, telefone, endereço e o link <strong>ens.edu.br</strong>.</li>"
            "<li>Se estiver tudo certo, clique em <strong>OK</strong> para salvar.</li>"
            "</ol>"
        )
        partes.append(_imagem_guia_html("image3.png", "Visualização da assinatura ENS no Outlook clássico", anexos))

        partes.append("<h4 style='margin-bottom:4px;'>Passo 4 - Definir como padrão e testar</h4>")
        partes.append(
            "<ol style='padding-left:20px;margin-top:4px;'>"
            "<li>Em <strong>Selecione as assinaturas padrão</strong>, escolha sua conta ENS.</li>"
            "<li>Defina <strong>Assinatura ENS</strong> para novas mensagens e respostas/encaminhamentos.</li>"
            "<li>Clique em <strong>OK</strong> e abra uma nova mensagem para testar.</li>"
            "<li>Se aparecer mais de uma assinatura, mantenha apenas a ENS.</li>"
            "</ol>"
        )
        partes.append(_imagem_guia_html("image4.png", "Configuração padrão da assinatura ENS", anexos))

        partes.append("<h3 style='color:#009db7;margin-top:24px;margin-bottom:8px;'>Parte 2 - Outlook Web / Novo Outlook</h3>")
        partes.append("<h4 style='margin-bottom:4px;'>Passo 5 - Configurar a assinatura no Outlook Web / Novo Outlook</h4>")
        partes.append(
            "<ol style='padding-left:20px;margin-top:4px;'>"
            "<li>Abra o Outlook no navegador ou o aplicativo <strong>Novo Outlook</strong>.</li>"
            "<li>Clique no ícone de engrenagem e depois em <strong>Ver todas as configurações do Outlook</strong>.</li>"
            "<li>Acesse <strong>Email &gt; Criar e responder</strong>.</li>"
            "<li>Clique em <strong>Nova assinatura</strong>, digite <strong>Assinatura ENS</strong> e cole (Ctrl+V) a assinatura.</li>"
            "<li>Em <strong>Escolher assinaturas padrão</strong>, selecione <strong>Assinatura ENS</strong> para novas mensagens e respostas.</li>"
            "<li>Salve e envie um e-mail de teste para confirmar.</li>"
            "</ol>"
        )
        partes.append(_imagem_guia_html("image5.png", "Assinatura ENS no Outlook Web / Novo Outlook", anexos))
        partes.append("<p style='margin-top:24px;'>Pronto! Sua assinatura corporativa ENS está configurada nos ambientes Outlook clássico e Outlook Web.</p>")
        partes.append("</div>")
        html = "".join(partes)
        _GUIA_ILUSTRADO_CACHE = (html, anexos)
        return html, [(cid, mime, data) for cid, mime, data in anexos]
    except Exception:
        _GUIA_ILUSTRADO_CACHE = ("", [])
        return None


def _montar_guia_textual_padrao() -> str:
    passos, dicas_adicionais = obter_passos_e_dicas_instalacao_outlook()
    partes: list[str] = []
    partes.append("<div style='font-family:\"Segoe UI\",Arial,sans-serif;color:#0b2630;line-height:1.55;font-size:14px;'>")
    partes.append("<h3 style='color:#009db7;'>Como instalar esta assinatura no Outlook</h3>")
    partes.append("<ol style='padding-left:20px;'>")
    for idx, passo in enumerate(passos, start=1):
        partes.append(f"<li>Passo {idx}: {passo}</li>")
    partes.append("</ol>")
    partes.append("<h4 style='color:#009db7;'>Dicas adicionais</h4>")
    partes.append("<ul style='padding-left:18px;'>")
    for dica in dicas_adicionais:
        partes.append(f"<li>{dica}</li>")
    partes.append("</ul>")
    partes.append("</div>")
    return "".join(partes)


def fonte_logo() -> str:
    """
    Retorna o logo codificado em data URI.
    Prioriza arquivos locais, com fallback para o payload base64 embarcado.
    """
    # Verifica nomes locais comuns (prioriza 'Logo.png')
    for candidate in (
        DIRETORIO_ICONES / "Logo.png",
        DIRETORIO_ICONES / "logo.png",
        DIRETORIO_ESTATICO / "logo.png",
    ):
        uri = _gerar_data_uri(candidate)
        if uri:
            return uri
    return _asset_fallback("logo")

def fonte_icone(name: str) -> str | None:
    """
    Retorna icone como data URI. Apenas arquivos locais no diretorio static/icons sao utilizados;
    URLs remotas nunca sao acionadas. Quando nao houver arquivo local, utiliza o fallback base64.
    """
    candidate = DIRETORIO_ICONES / f"{name}.png"
    uri = _gerar_data_uri(candidate)
    if uri:
        return uri
    # tenta fallback com nome capitalizado
    candidate = DIRETORIO_ICONES / f"{name.capitalize()}.png"
    uri = _gerar_data_uri(candidate)
    if uri:
        return uri
    candidate = DIRETORIO_ESTATICO / f"{name}.png"
    uri = _gerar_data_uri(candidate)
    if uri:
        return uri
    return _asset_fallback(name)

# ----- HTML da assinatura (final = previa) -----
COR_ASSINATURA_HEX = str(SIGNATURE_MODEL_SPEC["primary_color_hex"])
COR_ASSINATURA_RGB = str(SIGNATURE_MODEL_SPEC["primary_color_rgb"])
COR_AVISO_HEX = str(SIGNATURE_MODEL_SPEC["disclaimer_color_hex"])
COR_AVISO_RGB = str(SIGNATURE_MODEL_SPEC["disclaimer_color_rgb"])


def _fmt_pt(valor: float) -> str:
    numero = float(valor)
    if numero.is_integer():
        return f"{int(numero)}pt"
    return f"{numero:.1f}pt"


def _fmt_px_from_pt(valor: float) -> str:
    """
    Converte pontos tipograficos para pixels CSS (96dpi), com arredondamento estavel.
    """
    px = float(valor) * (96.0 / 72.0)
    if abs(px - round(px)) < 0.01:
        return f"{int(round(px))}px"
    texto = f"{px:.2f}".rstrip("0").rstrip(".")
    return f"{texto}px"


def _estilos_cor_primaria() -> list[str]:
    return [
        f"color:#{COR_ASSINATURA_HEX}",
        f"color:{COR_ASSINATURA_RGB}",
        f"color:#{COR_ASSINATURA_HEX} !important",
        f"-webkit-text-fill-color:#{COR_ASSINATURA_HEX}",
        f"mso-style-textfill-fill-color:#{COR_ASSINATURA_HEX}",
    ]


def _estilos_cor_aviso() -> list[str]:
    return [
        f"color:#{COR_AVISO_HEX}",
        f"color:{COR_AVISO_RGB}",
        f"color:#{COR_AVISO_HEX} !important",
        f"-webkit-text-fill-color:#{COR_AVISO_HEX}",
        f"mso-style-textfill-fill-color:#{COR_AVISO_HEX}",
    ]


def _bloco_aviso_html(rotulo: str, texto: str) -> str:
    disclaimer_pt = float(SIGNATURE_MODEL_SPEC["font_sizes_pt"]["disclaimer"])
    tamanho_css = _fmt_pt(disclaimer_pt)
    estilo_base = [
        "font-family:Verdana,Arial,sans-serif",
        f"font-size:{tamanho_css}",
        "line-height:105%",
        "mso-line-height-rule:exactly",
        *_estilos_cor_aviso(),
    ]
    estilo_str = ";".join(estilo_base)
    return (
        "<p style='margin:0;text-align:justify;"
        f"{estilo_str}"
        "'>"
        "<span style='display:inline;"
        f"{estilo_str}"
        "'>"
        f"<font face='Verdana, Arial, sans-serif' color='#{COR_AVISO_HEX}' "
        f"size='{map_pt_to_html_font_size(disclaimer_pt)}' style='{estilo_str}'>"
        f"<b style='{';'.join(_estilos_cor_aviso())}'>{rotulo}</b>: {texto}"
        "</font></span></p>"
    )


AVISO_PADRAO_HTML = (
    _bloco_aviso_html("IMPORTANTE", AVISO_TEXTO_PT)
    + "<br><br>"
    + _bloco_aviso_html("IMPORTANT", AVISO_TEXTO_EN)
)


def aviso_html() -> str:
    return AVISO_PADRAO_HTML


aviso_html_pt = aviso_html_en = aviso_html_bilingue = aviso_html


def _extrair_partes_aviso_docx() -> list[tuple[str, str]]:
    """
    Converte o HTML do aviso em pares (destaque, complemento) para aplicar estilos no Word.
    """
    return [
        ("IMPORTANTE", AVISO_TEXTO_PT),
        ("IMPORTANT", AVISO_TEXTO_EN),
    ]


def _assert_html_assinatura_canonica(html: str) -> None:
    if not isinstance(html, str) or not html.strip():
        raise ValueError("HTML canônico vazio.")
    faltando = [token for token in CANONICAL_HTML_REQUIRED_MARKERS if token not in html]
    if faltando:
        amostra = ", ".join(faltando[:6])
        raise ValueError(f"HTML canônico inválido. Tokens ausentes: {amostra}")

def linha_icones_html() -> str:
    links = []
    for nome, href in ICONE_LINKS_SOCIAIS.items():
        src = fonte_icone(nome)
        if not src:
            continue
        links.append(
            f"<a href='{href}' style='text-decoration:none;border:0;display:inline-block;'>"
            f"<img src='{src}' width='29' height='29' "
            "style='border:0;width:29px;height:29px;display:inline-block;' "
            f"alt='{nome.title()}'></a>"
        )
    if not links:
        return "&nbsp;"
    return "&nbsp;&nbsp;".join(links)

def _resolver_contexto_assinatura(dados: dict) -> dict[str, object]:
    """Extrai campos normalizados para montar HTML, texto e pacotes especiais."""
    display_name = (
        dados.get("nome_exibicao")
        or dados.get("campo_assinatura")
        or dados.get("name")
        or ""
    ).strip()
    department = (dados.get("department") or "").strip()
    role = (dados.get("role") or "").strip()
    phone = (dados.get("phone") or "").strip()
    eh_atn = eh_colaborador_atn(dados)
    local_descricao = (dados.get("local_descricao") or "").strip()
    endereco_informado_raw = dados.get("endereco") or dados.get("address")
    endereco_informado = str(endereco_informado_raw).strip() if endereco_informado_raw else ""
    tipo_inferido = inferir_tipo_endereco(dados)
    addr_type = (dados.get("address_type") or tipo_inferido or "rio").strip().lower()

    if eh_atn:
        if endereco_informado:
            address_line = endereco_informado
        elif addr_type == "sp":
            address_line = ENDERECO_SP
        elif addr_type in ("home", "homeoffice", "home-office"):
            address_line = ENDERECO_HOME
        else:
            address_line = ""
    else:
        address_line = normalizar_endereco_rio(endereco_informado, local_descricao)
        if not address_line:
            if endereco_informado:
                address_line = endereco_informado
            elif addr_type == "sp":
                address_line = ENDERECO_SP
            elif addr_type in ("home", "homeoffice", "home-office"):
                address_line = ENDERECO_HOME
            else:
                address_line = normalizar_endereco_rio(ENDERECO_RIO) or ENDERECO_RIO

    is_home_office, home_office_label = resolver_home_office(dados, tipo_inferido, address_line)
    if is_home_office:
        home_office_label = home_office_label or (ENDERECO_HOME or "").strip() or "Home Office"
        address_line = ""
    else:
        home_office_label = ""

    return {
        "name": display_name,
        "title": (department or role).strip(),
        "phone": phone,
        "address_line": (address_line or "").strip(),
        "is_home_office": is_home_office,
        "home_office_label": home_office_label,
    }


def assinatura_html_precisa(person: dict) -> str:
    """Layout ajustado pixel a pixel para reproduzir o modelo fornecido."""
    dados = person if isinstance(person, dict) else dict(person)
    contexto = _resolver_contexto_assinatura(dados)
    tamanhos = dict(SIGNATURE_MODEL_SPEC["font_sizes_pt"])
    largura_tabela = int(SIGNATURE_MODEL_SPEC["table_width_px"])
    cel_logo, cel_espacador, cel_texto = SIGNATURE_MODEL_SPEC["cell_widths_px"]
    altura_linha_principal = _fmt_pt(float(SIGNATURE_MODEL_SPEC["row_main_height_pt"]))
    altura_linha_espaco = _fmt_pt(float(SIGNATURE_MODEL_SPEC["row_spacer_height_pt"]))
    padding_lateral = _fmt_pt(float(SIGNATURE_MODEL_SPEC["content_padding_pt"]))
    logo_w, logo_h = SIGNATURE_MODEL_SPEC["logo_size_px"]

    def _linha_padrao(texto: str, *, tamanho_pt: float, bold: bool = False, italic: bool = False) -> str:
        tamanho_css = _fmt_pt(float(tamanho_pt))
        estilo_p = [
            "margin:0",
            "line-height:115%",
            "mso-line-height-rule:exactly",
            "font-family:Verdana,Arial,sans-serif",
            *_estilos_cor_primaria(),
            f"font-size:{tamanho_css}",
        ]
        estilo_span = [
            "display:inline",
            "font-family:Verdana,Arial,sans-serif",
            "line-height:115%",
            *_estilos_cor_primaria(),
            f"font-size:{tamanho_css}",
        ]
        if bold:
            estilo_p.append("font-weight:bold")
            estilo_span.append("font-weight:bold")
        if italic:
            estilo_p.append("font-style:italic")
            estilo_span.append("font-style:italic")
        conteudo_bruto = (texto or "").strip()
        conteudo = escape(conteudo_bruto) if conteudo_bruto else "&nbsp;"
        tamanho_html = map_pt_to_html_font_size(float(tamanho_pt))
        estilo_str = ";".join(estilo_span)
        html_interno = (
            f"<span style=\"{estilo_str}\">"
            f"<font face='Verdana, Arial, sans-serif' color='#{COR_ASSINATURA_HEX}' "
            f"size='{tamanho_html}' style=\"{estilo_str}\">{conteudo}</font>"
            "</span>"
        )
        if bold:
            html_interno = (
                f"<b style='font-weight:bold;{';'.join(_estilos_cor_primaria())}'>"
                f"{html_interno}</b>"
            )
        if italic:
            html_interno = (
                f"<i style='font-style:italic;{';'.join(_estilos_cor_primaria())}'>"
                f"{html_interno}</i>"
            )
        return f"<p style=\"{';'.join(estilo_p)}\">{html_interno}</p>"

    def _linha_site_e_icones(icones_inline: str) -> str:
        tamanho_site = _fmt_pt(float(tamanhos["site"]))
        estilos_base = [
            "font-family:Verdana,Arial,sans-serif",
            f"font-size:{tamanho_site}",
            "font-weight:bold",
            *_estilos_cor_primaria(),
        ]
        estilo_p = [
            "margin:0",
            "line-height:115%",
            "mso-line-height-rule:exactly",
            *estilos_base,
        ]
        estilo_span = ["display:inline", "line-height:115%", *estilos_base]
        link_style = ";".join(
            [
                *estilos_base,
                "text-decoration:none",
                "text-underline:none",
            ]
        )
        site_texto = (
            f"<span style='{';'.join(estilo_span)}'>"
            f"<font face='Verdana, Arial, sans-serif' color='#{COR_ASSINATURA_HEX}' "
            f"size='{map_pt_to_html_font_size(float(tamanhos['site']))}' style='{';'.join(estilo_span)}'>"
            "ens.edu.br"
            "</font></span>"
        )
        site_link = (
            "<a href='https://ens.edu.br' style='{style}'>"
            "{texto}"
            "</a>"
        ).format(style=link_style, texto=site_texto)
        return (
            f"<p style='{';'.join(estilo_p)}'>"
            f"<span style='{';'.join(estilo_span)}'>{site_link}</span><br><br>{icones_inline}"
            "</p>"
        )

    logo_src = fonte_logo()
    if logo_src:
        logo_html = (
            "<a href='https://ens.edu.br' style='text-decoration:none;border:0;'>"
            f"<img src='{logo_src}' width='{logo_w}' height='{logo_h}' "
            f"style='display:block;border:0;width:{logo_w}px;height:{logo_h}px' alt='Logotipo ENS'></a>"
        )
    else:
        logo_html = "&nbsp;"

    linhas_texto: list[str] = []
    linhas_texto.append(_linha_padrao(contexto.get("name", ""), tamanho_pt=float(tamanhos["name"]), bold=True))
    linhas_texto.append(_linha_padrao(contexto.get("title", ""), tamanho_pt=float(tamanhos["title"]), italic=True))
    if contexto.get("is_home_office") and contexto.get("home_office_label"):
        linhas_texto.append(_linha_padrao(contexto.get("home_office_label", ""), tamanho_pt=float(tamanhos["title"]), italic=True))
    linhas_texto.append(_linha_padrao(contexto.get("phone", ""), tamanho_pt=float(tamanhos["phone"])))
    linhas_texto.append(_linha_padrao("", tamanho_pt=float(tamanhos["title"])))
    linhas_texto.append(_linha_padrao(contexto.get("address_line", ""), tamanho_pt=float(tamanhos["address"])))
    linhas_texto.append(_linha_site_e_icones(linha_icones_html()))

    corpo_texto = "".join(linhas_texto)

    tabela_style = (
        "border-collapse:collapse;"
        f"width:{largura_tabela}px;min-width:{largura_tabela}px;max-width:{largura_tabela}px;"
        "font-family:Verdana,Arial,sans-serif;"
        f"{';'.join(_estilos_cor_primaria())};"
    )

    html = (
        f"<table role='presentation' cellpadding='0' cellspacing='0' border='0' width='{largura_tabela}' "
        f"style='{tabela_style}'>"
        f"<tr style='height:{altura_linha_principal}'>"
        f"<td valign='middle' width='{cel_logo}' style='width:113.35pt;padding:0 {padding_lateral} 0 {padding_lateral};height:{altura_linha_principal};vertical-align:middle;'>"
        f"{logo_html}"
        "</td>"
        f"<td width='{cel_espacador}' valign='top' style='width:14.2pt;padding:0;font-size:1pt;line-height:1;'>"
        "&nbsp;"
        "</td>"
        f"<td width='{cel_texto}' valign='top' style='width:410.55pt;padding:0 {padding_lateral} 0 {padding_lateral};height:{altura_linha_principal};font-family:Verdana,Arial,sans-serif;{';'.join(_estilos_cor_primaria())};'>"
        f"{corpo_texto}"
        "</td>"
        "</tr>"
        "<tr>"
        f"<td colspan='3' style='height:{altura_linha_espaco};padding:0 {padding_lateral} 0 {padding_lateral};font-size:1pt;line-height:1;'>"
        "&nbsp;"
        "</td>"
        "</tr>"
        "<tr>"
        "<td colspan='3' "
        f"style='border-top:0;padding:8pt {padding_lateral} 0 {padding_lateral};"
        "font-family:Verdana,Arial,sans-serif;"
        f"font-size:{_fmt_pt(float(tamanhos['disclaimer']))};line-height:105%;"
        "mso-line-height-rule:exactly;"
        f"{';'.join(_estilos_cor_aviso())};text-align:justify;'>"
        f"{aviso_html_bilingue()}"
        "</td>"
        "</tr>"
        "</table>"
    )
    _assert_html_assinatura_canonica(html)
    return html

def assinatura_html_outlook_new(person: dict) -> str:
    """
    Gera HTML dedicado para Outlook New/Web priorizando compatibilidade de colagem.
    Estrutura em tabelas e estilos inline por celula para reduzir perdas apos sanitizacao.
    """
    dados = person if isinstance(person, dict) else dict(person)
    contexto = _resolver_contexto_assinatura(dados)
    tamanhos = dict(SIGNATURE_MODEL_SPEC["font_sizes_pt"])
    largura_tabela = int(SIGNATURE_MODEL_SPEC["table_width_px"])
    cel_logo, cel_espacador, cel_texto = SIGNATURE_MODEL_SPEC["cell_widths_px"]
    padding_lateral = _fmt_pt(float(SIGNATURE_MODEL_SPEC["content_padding_pt"]))
    logo_w, logo_h = SIGNATURE_MODEL_SPEC["logo_size_px"]

    def _texto_style(
        tamanho_pt: float,
        *,
        cor_hex: str,
        bold: bool = False,
        italic: bool = False,
        line_height: str = "1.15",
    ) -> str:
        tamanho_px = _fmt_px_from_pt(float(tamanho_pt))
        partes = [
            "font-family:Verdana,Arial,sans-serif !important",
            f"font-size:{tamanho_px} !important",
            f"line-height:{line_height} !important",
            f"color:#{cor_hex} !important",
            "mso-line-height-rule:exactly",
        ]
        if bold:
            partes.append("font-weight:700 !important")
        if italic:
            partes.append("font-style:italic !important")
        return ";".join(partes)

    def _linha_texto(
        texto: str,
        *,
        tamanho_pt: float,
        bold: bool = False,
        italic: bool = False,
        padding_top: str = "0",
    ) -> str:
        conteudo = escape((texto or "").strip()) or "&nbsp;"
        return (
            "<tr>"
            f"<td style='padding:{padding_top} 0 0 0;vertical-align:top;{_texto_style(tamanho_pt, cor_hex=COR_ASSINATURA_HEX, bold=bold, italic=italic)}'>"
            f"{conteudo}"
            "</td>"
            "</tr>"
        )

    def _linha_site_e_icones(icones_inline: str) -> str:
        estilo_site = _texto_style(float(tamanhos["site"]), cor_hex=COR_ASSINATURA_HEX, bold=True)
        return (
            "<tr>"
            f"<td style='padding:0;vertical-align:top;{estilo_site}'>"
            f"<a href='https://ens.edu.br' style='{estilo_site};text-decoration:none !important;'>ens.edu.br</a>"
            "</td>"
            "</tr>"
            "<tr>"
            "<td style='padding:8px 0 0 0;vertical-align:top;line-height:1;'>"
            f"{icones_inline}"
            "</td>"
            "</tr>"
        )

    def _linha_aviso(rotulo: str, texto: str, *, margin_top_px: int = 0) -> str:
        margem = "0" if margin_top_px <= 0 else f"{margin_top_px}px"
        estilo = _texto_style(
            float(tamanhos["disclaimer"]),
            cor_hex=COR_AVISO_HEX,
            line_height="1.05",
        )
        return (
            "<tr>"
            f"<td style='padding:{margem} 0 0 0;text-align:justify;vertical-align:top;{estilo}'>"
            f"<strong style='font-weight:700 !important'>{escape(rotulo)}:</strong> {escape(texto)}"
            "</td>"
            "</tr>"
        )

    logo_src = fonte_logo()
    if logo_src:
        logo_html = (
            "<a href='https://ens.edu.br' style='text-decoration:none;border:0;display:inline-block;'>"
            f"<img src='{logo_src}' width='{logo_w}' height='{logo_h}' "
            f"style='display:block;border:0;width:{logo_w}px;height:{logo_h}px' alt='Logotipo ENS'>"
            "</a>"
        )
    else:
        logo_html = "&nbsp;"

    linhas_texto: list[str] = []
    linhas_texto.append(_linha_texto(contexto.get("name", ""), tamanho_pt=float(tamanhos["name"]), bold=True))
    linhas_texto.append(_linha_texto(contexto.get("title", ""), tamanho_pt=float(tamanhos["title"]), italic=True))
    if contexto.get("is_home_office") and contexto.get("home_office_label"):
        linhas_texto.append(_linha_texto(contexto.get("home_office_label", ""), tamanho_pt=float(tamanhos["title"]), italic=True))
    linhas_texto.append(_linha_texto(contexto.get("phone", ""), tamanho_pt=float(tamanhos["phone"])))
    linhas_texto.append(_linha_texto("", tamanho_pt=float(tamanhos["title"])))
    linhas_texto.append(_linha_texto(contexto.get("address_line", ""), tamanho_pt=float(tamanhos["address"])))
    linhas_texto.append(_linha_site_e_icones(linha_icones_html()))

    aviso_pt = _linha_aviso("IMPORTANTE", AVISO_TEXTO_PT, margin_top_px=0)
    aviso_en = _linha_aviso("IMPORTANT", AVISO_TEXTO_EN, margin_top_px=8)

    html = (
        f"<table role='presentation' cellpadding='0' cellspacing='0' border='0' width='{largura_tabela}' "
        f"style='border-collapse:collapse;width:{largura_tabela}px;min-width:{largura_tabela}px;max-width:{largura_tabela}px;background:#ffffff;'>"
        "<tr>"
        f"<td valign='middle' width='{cel_logo}' style='width:{cel_logo}px;padding:0 {padding_lateral} 0 {padding_lateral};vertical-align:middle;'>"
        f"{logo_html}"
        "</td>"
        f"<td valign='top' width='{cel_espacador}' style='width:{cel_espacador}px;padding:0;font-size:1px;line-height:1;'>&nbsp;</td>"
        f"<td valign='top' width='{cel_texto}' style='width:{cel_texto}px;padding:0 {padding_lateral} 0 {padding_lateral};vertical-align:top;'>"
        "<table role='presentation' cellpadding='0' cellspacing='0' border='0' width='100%' style='border-collapse:collapse;width:100%;'>"
        f"{''.join(linhas_texto)}"
        "</table>"
        "</td>"
        "</tr>"
        "<tr>"
        f"<td colspan='3' style='height:8.8pt;padding:0 {padding_lateral} 0 {padding_lateral};font-size:1px;line-height:1;'>&nbsp;</td>"
        "</tr>"
        "<tr>"
        f"<td colspan='3' style='padding:8pt {padding_lateral} 0 {padding_lateral};'>"
        "<table role='presentation' cellpadding='0' cellspacing='0' border='0' width='100%' style='border-collapse:collapse;width:100%;'>"
        f"{aviso_pt}{aviso_en}"
        "</table>"
        "</td>"
        "</tr>"
        "</table>"
    )
    return html

def html_para_texto(snippet: str) -> str:
    if not snippet:
        return ""
    texto = re.sub(r"(?i)<br\s*/?>", "\n", snippet)
    texto = re.sub(r"<[^>]+>", "", texto)
    texto = unescape(texto)
    texto = re.sub(r"\r?\n", "\n", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()

def _gerar_nome_base_assinatura(person: dict) -> str:
    departamento = (person.get("department") or "").strip()
    nome = (person.get("name") or "").strip()
    bruto = f"{departamento}_{nome}" if departamento else nome
    if not bruto:
        bruto = ASSINATURA_NOME_PADRAO
    texto = unicodedata.normalize("NFKD", bruto)
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    texto = re.sub(r"[^A-Za-z0-9]+", "_", texto)
    texto = texto.strip("_") or ASSINATURA_NOME_PADRAO
    return texto[:64]

def _carregar_bytes_asset(candidatos: list[Path], fallback_key: str) -> bytes | None:
    for candidate in candidatos:
        try:
            if candidate.exists():
                return candidate.read_bytes()
        except Exception:
            continue
    return _extrair_bytes_data_uri(_asset_fallback(fallback_key))

def _carregar_logo_bytes() -> bytes | None:
    candidatos = [
        DIRETORIO_ICONES / "Logo.png",
        DIRETORIO_ICONES / "logo.png",
        DIRETORIO_ESTATICO / "logo.png",
    ]
    return _carregar_bytes_asset(candidatos, "logo")

def _carregar_icone_bytes(nome: str) -> bytes | None:
    candidatos = [
        DIRETORIO_ICONES / f"{nome}.png",
        DIRETORIO_ICONES / f"{nome.capitalize()}.png",
        DIRETORIO_ESTATICO / f"{nome}.png",
    ]
    return _carregar_bytes_asset(candidatos, nome.lower())

def gerar_documento_word_assinatura(person: dict) -> io.BytesIO:
    if Document is None:
        raise RuntimeError("A biblioteca python-docx nao esta instalada.")
    if not ARQUIVO_MODELO_WORD.exists():
        raise RuntimeError(
            "O arquivo-modelo da assinatura (ASSINATURAS DE E-MAIL ... v21.23) nao foi encontrado em static/."
        )

    dados = person if isinstance(person, dict) else dict(person)
    contexto = _resolver_contexto_assinatura(dados)

    doc = Document(str(ARQUIVO_MODELO_WORD))
    tabela = doc.tables[0]
    tabela_elemento = tabela._tbl

    corpo = doc.element.body
    for child in list(corpo):
        if child is tabela_elemento:
            continue
        if child.tag.endswith("sectPr"):
            continue
        corpo.remove(child)

    cel_texto = tabela.rows[0].cells[2]
    paragrafos = list(cel_texto.paragraphs)
    if len(paragrafos) < 5:
        raise RuntimeError("Estrutura inesperada no modelo da assinatura Word.")

    def _run_principal(paragraph):
        if paragraph.runs:
            run = paragraph.runs[0]
            for extra in paragraph.runs[1:]:
                paragraph._element.remove(extra._element)
        else:
            run = paragraph.add_run()
        return run

    def _definir_texto(paragraph, texto: str | None, *, fallback_nbsp: bool = True):
        valor = (texto or "").strip()
        if not valor and fallback_nbsp:
            valor = "\u00A0"
        run = _run_principal(paragraph)
        run.text = valor
        return run

    p_nome, p_titulo, p_telefone = paragrafos[:3]
    p_linha_vazia = paragrafos[3]
    p_endereco = paragrafos[4]

    _definir_texto(p_nome, contexto.get("name") or ASSINATURA_NOME_PADRAO)
    _definir_texto(p_titulo, contexto.get("title"))

    if contexto.get("is_home_office") and contexto.get("home_office_label"):
        par_home = p_telefone.insert_paragraph_before("")
        par_home.style = p_titulo.style
        run_home = _definir_texto(par_home, contexto.get("home_office_label"))
        run_home.font.italic = True
    _definir_texto(p_telefone, contexto.get("phone"))
    _definir_texto(p_linha_vazia, "")
    if contexto.get("is_home_office"):
        _definir_texto(p_endereco, "")
    else:
        _definir_texto(p_endereco, contexto.get("address_line"))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def eh_admin():
    return session.get("eh_admin") is True

bp_admin = Blueprint("admin", __name__)

# Garante autenticacao em todas as rotas do admin, redirecionando para login quando necessario
@bp_admin.before_app_request
def _require_admin_auth():
    # Atua apenas em rotas do blueprint admin
    if request.blueprint != "admin":
        return
    # Permite acesso sem login apenas a login e rota de saude
    if request.endpoint in ("admin.login", "admin.admin_saude"):
        return
    if not eh_admin():
        return redirect(url_for("admin.login"))

@bp_admin.route("/logout", endpoint="logout")
def sair():
    session.clear()
    return redirect(url_for("admin.login"))

@bp_admin.get("/saude", endpoint="admin_saude")
def verificar_saude_admin():
    return jsonify({"status":"ok"})
bp_publico = Blueprint("public", __name__)

# ------------------------------------------------------------------
# Funcoes auxiliares e rotas do OAuth2 Microsoft

def construir_aplicativo_msal():
    """
    Constroi um ConfidentialClientApplication do MSAL quando ID e segredo estao configurados.
    Retorna None se msal nao estiver instalado ou se faltar configuracao. A autoridade
    padrao usa o tenant 'common' quando ENS_MS_TENANT_ID nao e informado.
    """
    if not (MS_ID_CLIENTE and MS_SEGREDO_CLIENTE) or ConfidentialClientApplication is None:
        return None
    authority = f"https://login.microsoftonline.com/{MS_ID_LOCATARIO}"
    return ConfidentialClientApplication(
        MS_ID_CLIENTE,
        authority=authority,
        client_credential=MS_SEGREDO_CLIENTE,
    )

@bp_publico.route("/auth/login", endpoint="ms_login")
def login_microsoft():
    """Inicia o fluxo de autorizacao OAuth2 com a Microsoft para envio de e-mails."""
    app_instance = construir_aplicativo_msal()
    if app_instance is None:
        # Mostra página amigável em vez de erro 500 quando a integração não está configurada
        return render_template(
            "oauth_not_configured.html",
            title="Integração Microsoft não configurada",
        ), 200
    # A URL de retorno precisa ser absoluta; url_for gera com _external=True
    redirect_uri = url_for("public.ms_callback", _external=True)
    # Solicita codigo de autorizacao com os escopos definidos
    auth_url = app_instance.get_authorization_request_url(
        scopes=MS_ESCOPOS.split(),
        redirect_uri=redirect_uri,
    )
    return redirect(auth_url)

@bp_publico.route("/auth/callback", endpoint="ms_callback")
def retorno_microsoft():
    """Processa o retorno OAuth2 da Microsoft e grava o token em sessao."""
    # Se configuracao ou msal estiverem ausentes, interrompe
    app_instance = construir_aplicativo_msal()
    if app_instance is None:
        # Mesmo comportamento da rota de login: página amigável em vez de HTTP 500
        return render_template(
            "oauth_not_configured.html",
            title="Integração Microsoft não configurada",
        ), 200
    code = request.args.get("code")
    if not code:
        return "Codigo de autorizacao ausente.", 400
    redirect_uri = url_for("public.ms_callback", _external=True)
    result = app_instance.acquire_token_by_authorization_code(
        code,
        scopes=MS_ESCOPOS.split(),
        redirect_uri=redirect_uri,
    )
    # Em caso de sucesso o resultado contem access_token e refresh_token
    if "access_token" in result:
        session["ms_token"] = result
        session.modified = True
        # Redireciona para a pagina inicial para que o usuario envie a assinatura
        return redirect(url_for("public.index"))
    # Em caso de erro inesperado inclui descricao quando disponivel
    error_desc = result.get("error_description") or result.get("error") or "Erro desconhecido"
    return f"Erro ao obter token: {error_desc}", 400

@bp_publico.route("/", methods=["GET","POST"], endpoint="index")
def pagina_inicial():
    erro_busca = ""
    termo_busca = ""
    colaborador = None
    colaborador_token = ""
    html_assinatura = None

    if request.method != "POST":
        session.pop("colaborador_lookup", None)
        session.pop("ultima_assinatura", None)
        session.pop("ultima_assinatura_token", None)
        session.modified = True

    lookup_cache = session.get("colaborador_lookup")
    if isinstance(lookup_cache, dict):
        registro_id = lookup_cache.get("registro_id")
        if registro_id:
            with obter_conexao() as conn:
                row = conn.execute("SELECT * FROM colaboradores WHERE id=?", (registro_id,)).fetchone()
            if row:
                colaborador_atual, dados_assinatura_atual = _preparar_colaborador_publico(dict(row))
                lookup_cache["colaborador"] = colaborador_atual
                lookup_cache["dados_assinatura"] = dados_assinatura_atual
                session.modified = True
        colaborador = lookup_cache.get("colaborador")
        colaborador_token = lookup_cache.get("token") or ""
        termo_busca = lookup_cache.get("termo") or termo_busca

    ultima_assinatura = session.get("ultima_assinatura")
    if isinstance(ultima_assinatura, dict):
        try:
            html_assinatura = assinatura_html_precisa(ultima_assinatura)
            _assert_html_assinatura_canonica(html_assinatura)
        except Exception:
            html_assinatura = None

    if request.method == "POST":
        action = (request.form.get("do") or "").strip().lower()
        if action == "buscar":
            q = (request.form.get("q") or "").strip()
            termo_busca = q
            colaborador = None
            colaborador_token = ""
            session.pop("colaborador_lookup", None)
            session.pop("ultima_assinatura", None)
            session.pop("ultima_assinatura_token", None)
            session.modified = True
            if not q:
                erro_busca = "Informe uma matricula ou e-mail para buscar."
            else:
                is_email = "@" in q and "." in q
                has_digit = any(ch.isdigit() for ch in q)
                has_space = any(ch.isspace() for ch in q)
                lookup_allowed = False
                if is_email:
                    if has_space:
                        erro_busca = "Remova espacos e informe o e-mail completo."
                    else:
                        lookup_allowed = True
                else:
                    if has_space or not has_digit:
                        erro_busca = "Busque usando uma matricula sem espacos ou um e-mail completo."
                    else:
                        lookup_allowed = True
                if lookup_allowed:
                    with obter_conexao() as conn:
                        row = conn.execute(
                            "SELECT * FROM colaboradores WHERE email=? COLLATE NOCASE OR matricula=? COLLATE NOCASE",
                            (q, q),
                        ).fetchone()
                    if row:
                        registro_dict = dict(row)
                        colaborador, dados_assinatura = _preparar_colaborador_publico(registro_dict)
                        colaborador_token = _armazenar_colaborador_lookup(colaborador, dados_assinatura, registro_id=registro_dict.get("id"))
                        cache = session.get("colaborador_lookup")
                        if isinstance(cache, dict):
                            cache["termo"] = q
                            session.modified = True
                        html_assinatura = None
                    else:
                        erro_busca = "Nenhum colaborador encontrado para essa matricula ou e-mail."
        elif action == "gerar":
            token = (request.form.get("colaborador_token") or "").strip()
            lookup_cache = session.get("colaborador_lookup") or {}
            if not token or token != lookup_cache.get("token"):
                erro_busca = "Localize um colaborador antes de gerar a assinatura oficial."
                colaborador = lookup_cache.get("colaborador")
                colaborador_token = lookup_cache.get("token") or ""
            else:
                dados_assinatura = lookup_cache.get("dados_assinatura") or {}
                if not dados_assinatura:
                    erro_busca = "Dados do colaborador indisponiveis. Realize a busca novamente."
                else:
                    try:
                        html_assinatura = assinatura_html_precisa(dados_assinatura)
                        _assert_html_assinatura_canonica(html_assinatura)
                    except Exception:
                        html_assinatura = None
                        erro_busca = "Nao foi possivel gerar a assinatura. Tente novamente."
                    else:
                        session["ultima_assinatura"] = dict(dados_assinatura)
                        token_preview = _armazenar_preview_html(html_assinatura)
                        if token_preview:
                            session["ultima_assinatura_token"] = token_preview
                        else:
                            session.pop("ultima_assinatura_token", None)
                        session.modified = True
                    colaborador = lookup_cache.get("colaborador")
                    colaborador_token = lookup_cache.get("token") or ""
        # demais acoes (ex.: legado) nao sao utilizadas no novo fluxo

    return render_template(
        "index.html",
        title=TITULO_APLICACAO,
        hero_image_url=IMAGEM_HEROI_PUBLICA if HERO_PUBLICO_ATIVO else None,
        hero_enabled=HERO_PUBLICO_ATIVO,
        termo_busca=termo_busca,
        erro_busca=erro_busca,
        colaborador=colaborador,
        colaborador_token=colaborador_token,
        html_assinatura=html_assinatura,
        USE_SMTP_EMAIL=USE_SMTP_EMAIL,
    )


def obter_passos_e_dicas_instalacao_outlook():
    """
    Retorna listas de passos e dicas para instalar a assinatura
    no Outlook Web/Novo Outlook e no Outlook classico (desktop).
    O texto foi pensado para ser autoexplicativo, acompanhando as telas
    que aparecem no Outlook, como nas ilustracoes de treinamento.
    """
    passos = [
        (
            "Acesse o Portal de Assinaturas ENS, preencha seus dados e clique em "
            "'Pre-visualizar / Exportar' para gerar a assinatura atualizada."
        ),
        (
            "Clique em 'Copiar assinatura' para colocar a assinatura completa na area de "
            "transferencia. Para colar diretamente no Outlook classico, prefira o download 'Word (Outlook classico)' "
            "e utilize o arquivo gerado como referencia dentro do Outlook."
        ),
        (
            "[Outlook classico para Windows] Abra o Outlook instalado no computador ENS, "
            "clique em 'Arquivo' no canto superior esquerdo e depois em 'Opcoes'."
        ),
        (
            "[Outlook classico para Windows] Na janela 'Opcoes do Outlook', clique em 'Email' "
            "no menu da esquerda e, na area 'Redigir mensagens', clique no botao 'Assinaturas...'."
        ),
        (
            "[Outlook classico para Windows] Na tela 'Assinaturas e papel de carta', clique em 'Novo', "
            "digite um nome como 'Assinatura ENS' e cole (Ctrl+V) a assinatura no quadro de edicao. "
            "Em 'Selecione as assinaturas padrao', escolha sua conta de e-mail e selecione 'Assinatura ENS' "
            "para 'Novas mensagens' e para 'Respostas/encaminhamentos'. Clique em 'OK' para salvar."
        ),
        (
            "[Outlook Web / Novo Outlook] Se voce usa o Outlook no navegador ou o aplicativo 'Novo Outlook', "
            "clique no icone de engrenagem no canto superior direito e depois em 'Ver todas as configuracoes do Outlook'."
        ),
        (
            "[Outlook Web / Novo Outlook] Acesse 'Email' > 'Criar e responder'. Clique em 'Nova assinatura', "
            "digite 'Assinatura ENS', cole (Ctrl+V) a assinatura no campo de edicao e, em 'Escolher assinaturas padrao', "
            "selecione 'Assinatura ENS' para novas mensagens e respostas/encaminhamentos. Clique em 'Salvar' e envie "
            "um e-mail de teste para conferir se a assinatura aparece igual ao modelo."
        ),
    ]

    dicas_adicionais = [
        (
            "Ao criar uma nova mensagem, verifique se aparece apenas uma assinatura. Se houver duas, "
            "volte nas configuracoes de assinaturas e remova a antiga, deixando apenas a 'Assinatura ENS' como padrao."
        ),
        (
            "Se a formatacao sofrer ajustes inesperados ao colar no Outlook, utilize a opcao "
            "'Manter formatacao de origem' sempre que possivel ou recorra ao download em Word para manter a "
            "assinatura id\u00eantica ao modelo oficial."
        ),
        (
            "No Outlook classico, se a assinatura parecer diferente, experimente abrir o arquivo "
            "'assinatura_ENS.html' ou o arquivo Word gerado pelo portal. Copie novamente a assinatura "
            "a partir dele e cole de novo no editor de assinaturas quando necessario."
        ),
        (
            "Apos salvar, reinicie o Outlook (feche e abra novamente) para garantir que a nova assinatura "
            "esteja disponivel em todas as janelas de composicao."
        ),
        (
            "Guarde o arquivo HTML gerado pelo portal: ele pode ser utilizado caso seja necessario "
            "reativar a assinatura em outro computador."
        ),
    ]

    return passos, dicas_adicionais

@bp_publico.get("/guia-assinatura.docx", endpoint="guia_assinatura_word")
def guia_assinatura_word():
    """
    Entrega o guia ilustrado em Word para que o usuario veja o passo a passo com imagens,
    tanto do Outlook classico quanto do Outlook Web / Novo Outlook.
    O arquivo deve estar em: assets/static/Guia_Assinatura_ENS_Ilustrado_v1.docx
    """
    if not GUIA_ILUSTRADO_DOC.exists():
        mensagem = (
            "O guia ilustrado nao esta disponivel nesta instalacao. "
            "Verifique se o arquivo 'Guia_Assinatura_ENS_Ilustrado_v1.docx' "
            "foi copiado para a pasta 'assets/static' do Portal de Assinaturas."
        )
        response = make_response(mensagem, 503)
        response.headers["Content-Type"] = "text/plain; charset=utf-8"
        return response

    return send_file(
        GUIA_ILUSTRADO_DOC,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name="Guia_Assinatura_ENS_Ilustrado.docx",
        as_attachment=True,
    )


def _obter_html_assinatura_canonica_da_sessao() -> str | None:
    dados = session.get("ultima_assinatura")
    token = session.get("ultima_assinatura_token")
    html = _ler_preview_html(token)
    if html:
        try:
            _assert_html_assinatura_canonica(html)
        except Exception:
            html = None
    if not html and isinstance(dados, dict):
        html = assinatura_html_precisa(dados)
        _assert_html_assinatura_canonica(html)
        token_preview = _armazenar_preview_html(html)
        if token_preview:
            session["ultima_assinatura_token"] = token_preview
        else:
            session.pop("ultima_assinatura_token", None)
        session.modified = True
    return html if isinstance(html, str) and html.strip() else None


@bp_publico.get("/assinatura/outlook-classico.html", endpoint="assinatura_outlook_html")
def baixar_outlook_classico_html():
    """
    Retorna a assinatura em HTML com recursos embutidos, ideal para copiar/colar no Outlook classico.
    Requer que o usuario tenha gerado uma assinatura nesta sessao.
    """
    html = _obter_html_assinatura_canonica_da_sessao()
    if not html:
        return "Gere sua assinatura antes de copiar para o Outlook classico.", 400
    return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})

@bp_publico.get("/assinatura/outlook-classico.docx", endpoint="assinatura_outlook_word")
def baixar_outlook_classico_word():
    """
    Entrega um arquivo .docx com a assinatura e imagens embutidas para facilitar o uso no Outlook classico.
    """
    dados = session.get("ultima_assinatura")
    if not isinstance(dados, dict):
        return "Gere sua assinatura antes de baixar o arquivo Word.", 400
    if Document is None:
        return (
            "A biblioteca python-docx nao esta instalada neste servidor. "
            "Instale 'python-docx' para habilitar o download em Word."
        ), 503
    base = _gerar_nome_base_assinatura(dados) or ASSINATURA_NOME_PADRAO
    arquivo = gerar_documento_word_assinatura(dados)
    return send_file(
        arquivo,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{base}_Outlook.docx",
    )


@bp_publico.get("/assinatura/outlook-new.html", endpoint="assinatura_outlook_new_html")
def baixar_outlook_new_html():
    """
    Retorna o HTML para colagem no Outlook New/Web.
    Usa um markup dedicado para reduzir sanitizacao durante a colagem.
    """
    dados = session.get("ultima_assinatura")
    if not isinstance(dados, dict):
        return "Gere sua assinatura antes de copiar para o Outlook New.", 400
    html = assinatura_html_outlook_new(dados)
    return make_response(html, 200, {"Content-Type": "text/html; charset=utf-8"})




def _enviar_email_smtp(destinatario: str, corpo_email: str, imagens_inline: list[tuple[str, str, bytes]] | None = None) -> tuple[bool, str]:
    """Envia o corpo de e-mail via SMTP usando as credenciais de homologação.

    Esta função é usada quando `USE_SMTP_EMAIL` está ativo. Ela monta uma
    mensagem `multipart/alternative` apenas com a versão HTML da assinatura
    e envia usando `SMTP_USER` como remetente.

    Retorna uma tupla `(sucesso, mensagem_erro)` para permitir tratamento
    amigável na rota HTTP.
    """
    try:
        msg = MIMEMultipart("related")
        msg["Subject"] = "Assinatura ENS - instalacao passo a passo"
        msg["From"] = SMTP_USER
        msg["To"] = destinatario

        corpo_alt = MIMEMultipart("alternative")
        msg.attach(corpo_alt)
        parte_html = MIMEText(corpo_email, "html", "utf-8")
        corpo_alt.attach(parte_html)
        if imagens_inline:
            for cid, mime, dados in imagens_inline:
                subtype = mime.split("/", 1)[-1] or "png"
                imagem = MIMEImage(dados, _subtype=subtype)
                imagem.add_header("Content-ID", f"<{cid}>")
                imagem.add_header("Content-Disposition", "inline", filename=f"{cid}.{subtype}")
                msg.attach(imagem)

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=120) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.sendmail(SMTP_USER, [destinatario], msg.as_string())
        return True, ""
    except Exception as exc:
        print(f"Erro ao enviar email via SMTP: {exc}")
        return False, str(exc)


def _montar_corpo_email_assinatura(html: str) -> tuple[str, list[tuple[str, str, bytes]]]:
    """
    Constrói o corpo do e-mail combinando a assinatura em HTML com o guia passo a passo.
    Retorna o HTML resultante e a lista de imagens inline (cid, mime, bytes).
    """
    assinatura_box = (
        "<div style='background:#f2fbfd;border:1px solid #9cd7e4;border-radius:16px;"
        "padding:24px;margin-bottom:24px;font-family:Segoe UI,Arial,sans-serif;color:#0b2630;'>"
        "<p style='margin:0 0 12px;font-size:15px;font-weight:600;'>"
        "Copie a assinatura oficial abaixo e cole no Outlook."
        "</p>"
        "<div style='background:#ffffff;border:1px dashed #9cd7e4;border-radius:12px;"
        "padding:18px;overflow:auto;'>"
        f"{html}"
        "</div>"
        "<p style='margin:16px 0 0;font-size:13px;color:#0f4a57;'>"
        "Selecione todo o bloco e pressione Ctrl+C (Windows) ou Cmd+C (Mac) para copiar sem perder a formatacao."
        "</p>"
        "</div>"
    )
    partes: list[str] = [
        assinatura_box,
        "<hr style='margin-top:24px;margin-bottom:16px;border:0;border-top:1px solid #cccccc;'>",
    ]
    anexos: list[tuple[str, str, bytes]] = []
    guia_ilustrado = _html_guia_assinatura_ilustrado()
    if guia_ilustrado:
        guia_html, imagens = guia_ilustrado
        partes.append(guia_html)
        anexos.extend(imagens)
    else:
        partes.append(_montar_guia_textual_padrao())
    partes.append(
        "<p style='font-family:Segoe UI,Arial,sans-serif;font-size:11px;color:#666666;'>"
        f"Guia gerado em {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}."
        "</p>"
    )
    return "".join(partes), anexos


@bp_publico.post("/enviar_assinatura", endpoint="enviar_assinatura")
def enviar_assinatura():
    """
    Envia a assinatura gerada para o destinatrio informado utilizando a Microsoft Graph API.
    Espera um payload JSON {"email": "destinatario@exemplo.com", "html": "<table>...</table>"}.
    Exige que o usurio tenha se autenticado via OAuth2; o token fica armazenado em
    session['ms_token']. Caso no haja token, retornamos 401 para que o cliente
    redirecione o usurio a /auth/login. Retorna JSON indicando sucesso ou erro.
    """
    try:
        data = request.get_json(force=True) or {}
    except Exception:
        return jsonify({"status": "error", "message": "JSON invalido"}), 400
    recipient = (data.get("email") or "").strip()
    html_payload = (data.get("html") or "").strip()
    if not recipient:
        return jsonify({"status": "error", "message": "Dados insuficientes"}), 400

    assinatura_session = session.get("ultima_assinatura")
    assinatura_html = None
    if isinstance(assinatura_session, dict):
        try:
            assinatura_html = assinatura_html_precisa(assinatura_session)
        except Exception:
            assinatura_html = None
    if not assinatura_html:
        assinatura_html = html_payload
    if not assinatura_html:
        return jsonify({"status": "error", "message": "Assinatura indisponivel"}), 400

    # Monta corpo do email com a assinatura + guia passo a passo em texto
    corpo_email, imagens_inline = _montar_corpo_email_assinatura(assinatura_html)

    # Modo hibrido: se USE_SMTP_EMAIL estiver ativo, envia via SMTP e nao usa Graph
    if USE_SMTP_EMAIL:
        sucesso, erro = _enviar_email_smtp(recipient, corpo_email, imagens_inline)
        if sucesso:
            return jsonify({"status": "ok", "mode": "smtp"})
        return jsonify(
            {
                "status": "error",
                "message": f"Erro ao enviar email via SMTP: {erro or 'falha nao especificada'}",
            }
        ), 500

    # Garante que exista token OAuth2 em sessao
    token_data = session.get("ms_token") or {}
    access_token = token_data.get("access_token")
    if not access_token:
        # Nao autenticado; orienta cliente a realizar login no Outlook
        return jsonify({"status": "error", "message": "Nao autenticado"}), 401
    # Monta payload JSON para Microsoft Graph
    message_payload = {
        "message": {
            "subject": "Assinatura ENS - instalacao passo a passo",
            "body": {
                "contentType": "HTML",
                "content": corpo_email,
            },
            "toRecipients": [
                {"emailAddress": {"address": recipient}}
            ],
        },
        # Salva em Itens Enviados para que o usuario tenha copia no Outlook
        "saveToSentItems": True,
    }
    if imagens_inline:
        anexos_graph = []
        for cid, mime, dados in imagens_inline:
            anexos_graph.append(
                {
                    "@odata.type": "#microsoft.graph.fileAttachment",
                    "name": f"{cid}.{mime.split('/')[-1]}",
                    "contentId": cid,
                    "isInline": True,
                    "contentType": mime,
                    "contentBytes": base64.b64encode(dados).decode("ascii"),
                }
            )
        message_payload["message"]["attachments"] = anexos_graph
    # Envia o e-mail via Microsoft Graph
    try:
        resp = requests.post(
            "https://graph.microsoft.com/v1.0/me/sendMail",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json=message_payload,
            timeout=30,
        )
        if resp.status_code >= 200 and resp.status_code < 300:
            return jsonify({"status": "ok"})
        # Se nao autorizado, remove token para que o cliente reautentique
        if resp.status_code == 401:
            session.pop("ms_token", None)
            return jsonify({"status": "error", "message": "Token expirado ou invalido"}), 401
        # Retorna erro generico para demais codigos de status
        print(f"Erro Graph API: {resp.status_code} {resp.text}")
        return jsonify({"status": "error", "message": "Erro ao enviar email via Graph"}), 500
    except requests.Timeout:
        print("Erro ao chamar Graph API: timeout")
        return jsonify({"status": "error", "message": "Timeout ao enviar email"}), 504
    except Exception as e:
        print(f"Erro ao chamar Graph API: {e}")
        return jsonify({"status": "error", "message": "Erro ao enviar email"}), 500



@bp_admin.route("/", endpoint="admin")
def painel_admin():
    """
    Exibe o painel administrativo agrupado por area (ENS x ATN), com busca restrita a matricula ou e-mail.
    """
    if not eh_admin():
        return redirect(url_for("admin.login"))
    with obter_conexao() as conn:
        registros_banco = conn.execute("SELECT * FROM colaboradores ORDER BY name").fetchall()
    termo_busca = (request.args.get("search") or "").strip()
    termo_normalizado = _normalizar_token(termo_busca)

    grupos: dict[str, list[dict]] = {
        "ENS": [],
        "ATN": [],
        "Sem Matrícula": [],
    }
    total_colaboradores = len(registros_banco)
    ativos_total = 0
    sem_matricula_total = 0
    home_office_total = 0
    atualizados_30_dias = 0
    limite_recente = datetime.datetime.now() - datetime.timedelta(days=30)

    for registro in registros_banco:
        dados = dict(registro)
        status_atual = (dados.get("status") or "on").strip().lower()
        if status_atual not in ("off", "0", "false"):
            ativos_total += 1
        if not (dados.get("matricula") or "").strip():
            sem_matricula_total += 1
        tipo_inferido = inferir_tipo_endereco(dados)
        eh_home, home_label = resolver_home_office(dados, tipo_inferido)
        if eh_home:
            dados["endereco"] = home_label
            home_office_total += 1
        else:
            endereco_normalizado = normalizar_endereco_rio(dados.get("endereco"), dados.get("local_descricao"))
            if endereco_normalizado:
                dados["endereco"] = endereco_normalizado
        atualizado_em = dados.get("updated_at") or ""
        if atualizado_em:
            try:
                instante = datetime.datetime.fromisoformat(atualizado_em)
            except ValueError:
                instante = None
            if instante and instante >= limite_recente:
                atualizados_30_dias += 1
        matricula_normalizada = _normalizar_token(dados.get("matricula"))
        email_normalizado = _normalizar_token(dados.get("email"))
        if termo_normalizado and termo_normalizado not in matricula_normalizada and termo_normalizado not in email_normalizado:
            continue
        dados["lookup"] = " ".join(token for token in (matricula_normalizada, email_normalizado) if token)
        area = classificar_area_por_matricula(dados.get("matricula"))
        dados["area"] = area
        grupos.setdefault(area, []).append(dados)

    for chave in ("ENS", "ATN", "Sem Matrícula"):
        grupos.setdefault(chave, [])

    grupos_para_renderizar = []
    for chave, titulo in (
        ("ENS", "Colaboradores ENS"),
        ("ATN", "Colaboradores ATN"),
        ("Sem Matrícula", "Sem matrícula informada"),
    ):
        itens = grupos.get(chave, [])
        if itens:
            grupos_para_renderizar.append({"key": chave, "title": titulo, "rows": itens})

    possui_registros = any(grupos.values())
    quantidade_importada = request.args.get("importados")
    percentual_ativos = int(round((ativos_total / total_colaboradores) * 100)) if total_colaboradores else 0
    percentual_home = int(round((home_office_total / total_colaboradores) * 100)) if total_colaboradores else 0
    painel_resumo = [
        {
            "key": "total",
            "title": "Colaboradores",
            "value": total_colaboradores,
            "meta": f"{ativos_total} ativos • {percentual_ativos}%",
            "icon": "bi-people-fill",
        },
        {
            "key": "home",
            "title": "Home office",
            "value": home_office_total,
            "meta": f"{percentual_home}% do total",
            "icon": "bi-laptop",
        },
        {
            "key": "sem_matricula",
            "title": "Sem matrícula",
            "value": sem_matricula_total,
            "meta": "Preencher para liberar assinatura",
            "icon": "bi-exclamation-circle",
        },
        {
            "key": "atualizados",
            "title": "Atualizados (30 dias)",
            "value": atualizados_30_dias,
            "meta": "Movimentações recentes",
            "icon": "bi-clock-history",
        },
    ]
    return render_template(
        "admin.html",
        title=TITULO_APLICACAO,
        grupos=grupos_para_renderizar,
        quantidade_importada=quantidade_importada,
        termo_busca=termo_busca,
        possui_registros=possui_registros,
        letras=list("ABCDEFGHIJKLMNOPQRSTUVWXYZ"),
        painel_resumo=painel_resumo,
    )
@bp_admin.post("/add", endpoint="admin_add")
def adicionar_colaborador_admin():
    if not eh_admin():
        abort(403)
    dados = coletar_dados_colaborador(request.form)
    ajustar_cargo_com_departamento(dados)
    nome = dados.get("name") or ""
    email = dados.get("email") or ""
    if not nome or not email:
        return redirect(url_for("admin.admin"))
    endereco_normalizado = normalizar_endereco_rio(dados.get("endereco"), dados.get("local_descricao"))
    if endereco_normalizado:
        dados["endereco"] = endereco_normalizado
    updated_at = datetime.datetime.now().isoformat(timespec="seconds")
    with obter_conexao() as conn:
        # Busca registro existente por e-mail ou matricula para preservar flags, senha e campos extras
        existing_rows = conn.execute(
            "SELECT * FROM colaboradores WHERE email=? OR matricula=?",
            (email, dados.get("matricula"))
        ).fetchall()
        existing_dict = _selecionar_registro_existente(existing_rows, email, nome) or {}
        existing_id = existing_dict.get("id")
        status = existing_dict.get("status") or "on"
        eh_admin_flag = existing_dict.get("eh_admin") or "no"
        if existing_dict:
            # Quando ja existe, preserva valores nao fornecidos explicitamente
            for campo in CAMPOS_COLABORADOR_EDITAVEIS:
                if dados.get(campo) is None:
                    dados[campo] = existing_dict.get(campo)
        password_hash_val = existing_dict.get("password_hash") if existing_dict else None
        must_change_val = existing_dict.get("must_change") if existing_dict else "no"
    # Insere ou substitui registro incluindo colunas estendidas
        conn.execute(
            "INSERT OR REPLACE INTO colaboradores (id, matricula, name, department, role, phone, email, status, updated_at, eh_admin, password_hash, must_change, "
            "posicao_organograma, first_name, last_name, nome_exibicao, diretoria, campo_assinatura, manager, telefone_ad, local_descricao, endereco, uf) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                existing_id,
                dados.get("matricula"),
                nome,
                dados.get("department"),
                dados.get("role"),
                dados.get("phone"),
                email,
                status,
                updated_at,
                eh_admin_flag,
                password_hash_val,
                must_change_val,
                dados.get("posicao_organograma"),
                dados.get("first_name"),
                dados.get("last_name"),
                dados.get("nome_exibicao"),
                dados.get("diretoria"),
                dados.get("campo_assinatura"),
                dados.get("manager"),
                dados.get("telefone_ad"),
                dados.get("local_descricao"),
                dados.get("endereco"),
                dados.get("uf"),
            )
        )
    return redirect(url_for("admin.admin"))

@bp_admin.post("/status/<int:cid>", endpoint="admin_toggle")
def alternar_status_colaborador(cid: int):
    if not eh_admin():
        abort(403)
    with obter_conexao() as conn:
        row = conn.execute("SELECT status FROM colaboradores WHERE id=?", (cid,)).fetchone()
        if not row: return redirect(url_for("admin.admin"))
        new_status = "off" if row["status"] == "on" else "on"
        conn.execute("UPDATE colaboradores SET status=?, updated_at=? WHERE id=?", (new_status, datetime.datetime.now().isoformat(timespec="seconds"), cid))
    return redirect(url_for("admin.admin"))

@bp_admin.post("/status/<int:cid>/off", endpoint="admin_disable")
def desligar_colaborador(cid: int):
    """
    Marca o colaborador como desligado ('status' = 'off') sem remover o registro.
    Disponibilizado para casos em que o admin deseja apenas invalidar a assinatura.
    """
    if not eh_admin():
        abort(403)
    timestamp = datetime.datetime.now().isoformat(timespec="seconds")
    with obter_conexao() as conn:
        conn.execute(
            "UPDATE colaboradores SET status=?, updated_at=? WHERE id=?",
            ("off", timestamp, cid),
        )
    return redirect(url_for("admin.admin_edit", cid=cid))

@bp_admin.post("/delete/<int:cid>", endpoint="admin_delete")
def excluir_colaborador(cid: int):
    if not eh_admin():
        abort(403)
    with obter_conexao() as conn:
        conn.execute("DELETE FROM colaboradores WHERE id=?", (cid,))
    return redirect(url_for("admin.admin"))

@bp_admin.post("/admin_role/<int:cid>", endpoint="admin_role")
def alternar_papel_admin(cid: int):
    if not eh_admin():
        abort(403)
    with obter_conexao() as conn:
        row = conn.execute("SELECT eh_admin FROM colaboradores WHERE id=?", (cid,)).fetchone()
        if not row:
            return redirect(url_for("admin.admin"))
        current = str(row["eh_admin"]).lower()
        new_val = "no" if current in ("yes","on","true","1") else "yes"
        conn.execute("UPDATE colaboradores SET eh_admin=?, updated_at=? WHERE id=?", (new_val, datetime.datetime.now().isoformat(timespec="seconds"), cid))
    return redirect(url_for("admin.admin"))

@bp_admin.post("/update/<int:cid>", endpoint="admin_update")
def atualizar_colaborador(cid: int):
    if not eh_admin():
        abort(403)
    new_department = (request.form.get("department") or "").strip()
    new_role = (request.form.get("role") or "").strip()
    if not new_department and not new_role:
        return redirect(url_for("admin.admin"))
    fields = []
    params = []
    if new_department:
        fields.append("department=?")
        params.append(new_department)
    if new_role:
        fields.append("role=?")
        params.append(new_role)
    fields.append("updated_at=?")
    params.append(datetime.datetime.now().isoformat(timespec="seconds"))
    params.append(cid)
    with obter_conexao() as conn:
        conn.execute(f"UPDATE colaboradores SET {' , '.join(fields)} WHERE id=?", params)
    return redirect(url_for("admin.admin"))

# ----------------------------------------------------------------------
# Editar colaborador
@bp_admin.route("/edit/<int:cid>", methods=["GET", "POST"], endpoint="admin_edit")
def editar_colaborador(cid: int):
    """
    Permitir a edicao completa dos dados de um colaborador. Quando acessado via
    GET, exibe um formulario preenchido com os valores atuais. Quando
    submetido via POST, atualiza todos os campos editaveis no banco de dados
    e redireciona de volta ao painel de administracao. Somente usuarios
    autenticados como administradores podem acessar esta rota.
    """
    if not eh_admin():
        abort(403)
    # Carregar registro existente
    with obter_conexao() as conn:
        existing = conn.execute("SELECT * FROM colaboradores WHERE id=?", (cid,)).fetchone()
    if not existing:
        abort(404)
    # Se for envio de formulario, processar atualizacoes
    if request.method == "POST":
        data = coletar_dados_colaborador(request.form)
        ajustar_cargo_com_departamento(data)
        normalized_address = normalizar_endereco_rio(data.get("endereco"), data.get("local_descricao"))
        if normalized_address:
            data["endereco"] = normalized_address
        # Executar a atualizacao; campos nao fornecidos permanecem iguais a
        # None ou string vazia conforme design do banco. Nao alteramos
        # status, eh_admin, password_hash ou must_change aqui.
        with obter_conexao() as conn:
            conn.execute(
                "UPDATE colaboradores SET matricula=?, name=?, department=?, role=?, phone=?, email=?, "
                "posicao_organograma=?, first_name=?, last_name=?, nome_exibicao=?, diretoria=?, "
                "campo_assinatura=?, manager=?, telefone_ad=?, local_descricao=?, endereco=?, uf=?, "
                "updated_at=? WHERE id=?",
                (
                    data.get("matricula"),
                    data.get("name"),
                    data.get("department"),
                    data.get("role"),
                    data.get("phone"),
                    data.get("email"),
                    data.get("posicao_organograma"),
                    data.get("first_name"),
                    data.get("last_name"),
                    data.get("nome_exibicao"),
                    data.get("diretoria"),
                    data.get("campo_assinatura"),
                    data.get("manager"),
                    data.get("telefone_ad"),
                    data.get("local_descricao"),
                    data.get("endereco"),
                    data.get("uf"),
                    datetime.datetime.now().isoformat(timespec="seconds"),
                    cid,
                ),
            )
        return redirect(url_for("admin.admin"))
    # Caso GET, exibir o formulario de edicao
    return render_template("admin_edit.html", title=TITULO_APLICACAO, col=dict(existing))


@bp_admin.get("/assinatura/<int:cid>", endpoint="visualizar_assinatura")
def visualizar_assinatura(cid: int):
    if not eh_admin():
        abort(403)
    with obter_conexao() as conn:
        row = conn.execute("SELECT * FROM colaboradores WHERE id=?", (cid,)).fetchone()
    if not row:
        abort(404)
    html = assinatura_html_precisa(dict(row))
    return make_response(html, 200, {"Content-Type":"text/html; charset=utf-8"})


@bp_admin.get("/assinatura/outlook-new/<int:cid>", endpoint="visualizar_assinatura_outlook_new")
def visualizar_assinatura_outlook_new(cid: int):
    if not eh_admin():
        abort(403)
    with obter_conexao() as conn:
        row = conn.execute("SELECT * FROM colaboradores WHERE id=?", (cid,)).fetchone()
    if not row:
        abort(404)
    html = assinatura_html_outlook_new(dict(row))
    return make_response(html, 200, {"Content-Type":"text/html; charset=utf-8"})

@bp_admin.get("/baixar/<int:cid>/<ext>", endpoint="baixar_assinatura")
def baixar_assinatura(cid: int, ext: str):
    if not eh_admin():
        abort(403)
    if ext not in ("html","htm","docx"):
        abort(400)
    with obter_conexao() as conn:
        row = conn.execute("SELECT * FROM colaboradores WHERE id=?", (cid,)).fetchone()
    if not row:
        abort(404)
    person = dict(row)
    base_nome = _gerar_nome_base_assinatura(person)
    if ext == "docx":
        if Document is None:
            return make_response(
                "A biblioteca python-docx nao esta instalada neste servidor.",
                503,
                {"Content-Type": "text/plain; charset=utf-8"},
            )
        arquivo = gerar_documento_word_assinatura(person)
        download_name = f"{base_nome}_Outlook.docx"
        return send_file(
            arquivo,
            as_attachment=True,
            download_name=download_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    html = assinatura_html_precisa(person).encode("utf-8")
    filename = f"{base_nome}.{ext}"
    return send_file(io.BytesIO(html), as_attachment=True, download_name=filename, mimetype="text/html; charset=utf-8")

# ----------------------------------------------------------------------
# Integracao de importacao/exportacao de planilha

@bp_admin.post("/import", endpoint="admin_import")
def importar_csv_admin():
    """Importar colaboradores a partir de uma planilha Excel ou CSV.
    Espera um campo de arquivo chamado 'file' contendo um .xlsx/.xls/.csv.
    A planilha deve possuir colunas compativeis com a planilha base, incluindo:
    - Posicao no Organograma
    - Matricula (Descricao)\n(AD)
    - Nome e Sobrenome ou Nome para exibicao\n(AD)
    - Diretoria
    - Departamento (Departamento)\n(AD)
    - Cargo (Cargo)\n(AD)
    - Campo Assinatura
    - Manager (Gerente)\n(AD)
    - Telefone\n(ASSINATURA)
    - Telefone\n(AD)
    - Local (Escritorio)\n(AD)
    - Endereco e UF
    """
    if not eh_admin():
        abort(403)
    file = request.files.get("file")
    if not file:
        return redirect(url_for("admin.admin"))
    # Tenta ler o arquivo enviado utilizando pandas
    try:
        filename = (file.filename or "").lower()
        if filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file, engine="openpyxl")
        else:
    # Assume CSV por padrao; pandas detecta delimitador automaticamente
            df = pd.read_csv(file)
    except Exception:
    # Falha ao interpretar arquivo; ignora importacao silenciosamente
        return redirect(url_for("admin.admin"))
    def _normalize_header(header: object) -> str:
        text = str(header or "")
        text = unicodedata.normalize("NFKD", text)
        text = "".join(ch for ch in text if not unicodedata.combining(ch))
        text = re.sub(r"\s+", " ", text).strip().lower()
        text = re.sub(r"[^a-z0-9]+", "_", text)
        return text.strip("_")

    header_lookup: dict[str, str] = {}
    for column in df.columns:
        normalized = _normalize_header(column)
        if normalized and normalized not in header_lookup:
            header_lookup[normalized] = column

    quantidade_importada = 0
    with obter_conexao() as conn:
        for _, row in df.iterrows():
    # Extrai valores com seguranca; usa string vazia como fallback
    # Funcoes auxiliares: converte para string, remove espacos e trata 'nan' como vazio
            def _clean(cell: object) -> str:
    # Converte para string (tratando float/NaN), aplica strip e trata 'nan' ou 'NaN' como vazio
                if cell is None:
                    return ""
                s = str(cell).strip()
                return "" if s.lower() == "nan" else s

            def _pick(*aliases: str) -> object:
                for alias in aliases:
                    normalized = _normalize_header(alias)
                    if not normalized:
                        continue
                    column_name = header_lookup.get(normalized)
                    if not column_name:
                        continue
                    return row.get(column_name, "")
                return ""

            pos_org = _clean(_pick("Posicao no Organograma", "Posicao do Organograma"))
    # Algumas colunas possuem quebra de linha no cabecalho; testa ambas as variacoes
            matricula = _clean(_pick("Matricula (Descricao)\n(AD)", "Matricula (Descricao) (AD)", "Matricula (Descricao)", "Matricula"))
            first_name = _clean(_pick("Nome", "Primeiro Nome"))
            last_name = _clean(_pick("Sobrenome", "Ultimo Nome"))
            nome_exib = _clean(_pick("Nome para exibicao\n(AD)", "Nome para exibicao (AD)", "Nome para exibicao"))
            diretoria = _clean(_pick("Diretoria"))
            dept = _clean(_pick("Departamento (Departamento)\n(AD)", "Departamento (Departamento) (AD)", "Departamento"))
            role = _clean(_pick("Cargo (Cargo)\n(AD)", "Cargo (Cargo) (AD)", "Cargo"))
            campo_ass = _clean(_pick("Campo Assinatura"))
            manager = _clean(_pick("Manager (Gerente)\n(AD)", "Manager (Gerente) (AD)", "Gerente"))
            email = _clean(_pick("E-mail", "Email"))
            phone = _clean(_pick("Telefone\n(ASSINATURA)", "Telefone (ASSINATURA)", "Telefone assinatura"))
            tel_ad = _clean(_pick("Telefone\n(AD)", "Telefone (AD)", "Telefone ad"))
            local_desc = _clean(_pick("Local (Escritorio)\n(AD)", "Local (Escritorio) (AD)", "Local (Escritorio)", "Local"))
            endereco_raw = _pick(
                "Endereco",
                "Endereco e UF",
                "Endereco (AD)",
                "Endereco assinatura",
                "Endere\u00e7o",
                "Endere\u00e7o e UF",
                "Endere\u00e7o (AD)",
                "Endere\u00e7o assinatura",
            )
            endereco = _clean(endereco_raw)
            endereco = normalizar_endereco_rio(endereco, local_desc) or endereco
            uf_val = _clean(_pick("UF", "Uf", "UF ", "Unidade Federativa"))
            if uf_val:
                uf_val = uf_val.upper()
            if not uf_val and endereco_raw:
                match_uf = re.search(r"\b([A-Z]{2})\s*$", _clean(endereco_raw))
                if match_uf:
                    uf_val = match_uf.group(1)
    # Determina nome de exibicao: prefere Nome para exibicao; caso contrario junta nome e sobrenome
            name = nome_exib if nome_exib else (first_name + (" " + last_name if last_name else ""))
    # Ignora linhas sem e-mail e sem matr\u00edcula
            if not email and not matricula:
                continue
    # Busca registro existente para preservar flags e senha
            existing_rows = conn.execute(
                "SELECT * FROM colaboradores WHERE email=? OR matricula=?",
                (email, matricula)
            ).fetchall()
            existing_dict = _selecionar_registro_existente(existing_rows, email, name) or {}
            existing_id = existing_dict.get("id")
            status_val = existing_dict.get("status") or "on"
            eh_admin_val = existing_dict.get("eh_admin") or "no"
            password_hash_val = existing_dict.get("password_hash") if existing_dict else None
            must_change_val = existing_dict.get("must_change") if existing_dict else "no"
    # Se a matricula recebida estiver vazia, mantem o valor existente (quando houver) para evitar perda do registro anterior.
            if matricula:
                matricula_val = matricula
            else:
                matricula_val = existing_dict.get("matricula")
            updated_at = datetime.datetime.now().isoformat(timespec="seconds")
            conn.execute(
                "INSERT OR REPLACE INTO colaboradores (id, matricula, name, department, role, phone, email, status, updated_at, eh_admin, password_hash, must_change, "
                "posicao_organograma, first_name, last_name, nome_exibicao, diretoria, campo_assinatura, manager, telefone_ad, local_descricao, endereco, uf) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (
                    existing_id,
                    matricula_val,
                    name,
                    dept,
                    role,
                    phone,
                    email,
                    status_val,
                    updated_at,
                    eh_admin_val,
                    password_hash_val,
                    must_change_val,
                    pos_org,
                    first_name,
                    last_name,
                    nome_exib,
                    diretoria,
                    campo_ass,
                    manager,
                    tel_ad,
                    local_desc,
                    endereco,
                    uf_val,
                )
            )
            quantidade_importada += 1
    return redirect(url_for("admin.admin", importados=quantidade_importada))


@bp_admin.get("/import_ad", endpoint="admin_import_ad")
def importar_dados_ad():
    """
    Importar colaboradores diretamente do Azure Active Directory usando a API
    Microsoft Graph. Esse endpoint utiliza autenticacao do tipo client
    credentials (apponly) para solicitar um token e ler todos os usuarios
    presentes na organizacao.  Apenas administradores podem acessalo.
    Apos a conclusao, redireciona de volta ao painel admin e exibe o numero
    de registros importados na mensagem de sucesso.
    """
    if not eh_admin():
        abort(403)
    try:
        total_importado = importar_usuarios_do_graph()
    except Exception:
        total_importado = 0
    return redirect(url_for("admin.admin", importados=total_importado))


@bp_admin.get("/export", endpoint="admin_export")
def exportar_csv_admin():
    """Exportar os colaboradores atuais para uma planilha Excel (.xlsx).
    A planilha gerada agora inclui todas as colunas presentes na tabela 'colaboradores'.
    """
    if not eh_admin():
        abort(403)
    with obter_conexao() as conn:
        rows = conn.execute("SELECT * FROM colaboradores ORDER BY name").fetchall()
        col_meta = conn.execute("PRAGMA table_info(colaboradores)").fetchall()
    column_names = [info["name"] for info in col_meta] if col_meta else list(rows[0].keys()) if rows else []
    records: list[dict[str, object]] = []
    for row in rows:
        record: dict[str, object] = {}
        for column in column_names:
            value = row[column] if column in row.keys() else None
            record[column] = "" if value is None else value
        records.append(record)
    df = pd.DataFrame(records, columns=column_names)
    # Escrever DataFrame em memoria
    output = io.BytesIO()
    df.to_excel(output, index=False, engine="openpyxl")
    output.seek(0)
    # Nome do arquivo com data/hora
    stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"colaboradores_atualizados_{stamp}.xlsx"
    return send_file(output, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@bp_admin.get("/export/zip", endpoint="export_zip")
def exportar_zip_assinaturas():
    if not eh_admin():
        abort(403)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        with obter_conexao() as conn:
            rows = conn.execute("SELECT * FROM colaboradores WHERE status='on' ORDER BY department, name").fetchall()
        for r in rows:
            person = dict(r)
            base = _gerar_nome_base_assinatura(person)
            html = assinatura_html_precisa(person).encode("utf-8")
            z.writestr(base + ".html", html)
            z.writestr(base + ".htm",  html)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="assinaturas_ENS.zip", mimetype="application/zip")

# ----- Gerenciamento de senhas do admin -----
@bp_admin.post("/set_password/<int:cid>", endpoint="admin_set_password")
def definir_senha_colaborador(cid: int):
    if not eh_admin():
        abort(403)
    newpass = (request.form.get("newpass") or "").strip()
    if len(newpass) < 6:
        return redirect(url_for("admin.admin"))
    ph = generate_password_hash(newpass)
    with obter_conexao() as conn:
        conn.execute(
            "UPDATE colaboradores SET password_hash=?, must_change='yes', updated_at=? WHERE id=?",
            (ph, datetime.datetime.now().isoformat(timespec="seconds"), cid),
        )
    return redirect(url_for("admin.admin"))

@bp_admin.route("/login", methods=["GET","POST"], endpoint="login")
def login_admin():
    if request.method == "POST":
        user = (request.form.get("user") or "").strip()
        pwd  = (request.form.get("pass") or "").strip()
    # Administrador mestre
        if user == USUARIO_ADMINISTRADOR and pwd == SENHA_ADMINISTRADOR:
            session.clear()
            session["eh_admin"] = True
            session["admin_email"] = None
            return redirect(url_for("admin.admin"))
    # Administrador colaborador com senha individual
        if "@" in user:
            with obter_conexao() as conn:
                row = conn.execute(
                    "SELECT email, status, eh_admin, password_hash, must_change FROM colaboradores WHERE email=?",
                    (user,),
                ).fetchone()
            if row and row["status"] == "on" and str(row["eh_admin"]).lower() in ("yes","on","true","1") and row["password_hash"]:
                if check_password_hash(row["password_hash"], pwd):
                    session.clear()
                    session["eh_admin"] = True
                    session["admin_email"] = row["email"]
                    if str(row["must_change"]).lower() in ("yes","on","true","1"):
                        return redirect(url_for("admin.change_password"))
                    return redirect(url_for("admin.admin"))
        return render_template("login.html", title=TITULO_APLICACAO, error="Credenciais invalidas.")
    return render_template("login.html", title=TITULO_APLICACAO)

@bp_admin.route("/change_password", methods=["GET","POST"], endpoint="change_password")
def alterar_senha_admin():
    if not eh_admin() or not session.get("admin_email"):
        return redirect(url_for("admin.login"))
    email = session.get("admin_email")
    if request.method == "POST":
        current = (request.form.get("current") or "").strip()
        new1 = (request.form.get("new1") or "").strip()
        new2 = (request.form.get("new2") or "").strip()
        if len(new1) < 6 or new1 != new2:
            return render_template("change_password.html", title=TITULO_APLICACAO, error="Senha invalida ou diferente.")
        with obter_conexao() as conn:
            row = conn.execute("SELECT password_hash FROM colaboradores WHERE email=?", (email,)).fetchone()
            if not row or not row["password_hash"] or not check_password_hash(row["password_hash"], current):
                return render_template("change_password.html", title=TITULO_APLICACAO, error="Senha atual incorreta.")
            conn.execute(
                "UPDATE colaboradores SET password_hash=?, must_change='no', updated_at=? WHERE email=?",
                (generate_password_hash(new1), datetime.datetime.now().isoformat(timespec="seconds"), email),
            )
        return redirect(url_for("admin.admin"))
    return render_template("change_password.html", title=TITULO_APLICACAO)


# Instancia padrao para servidores WSGI (ex.: gunicorn, IIS, waitress-serve)
app = criar_aplicativo(os.getenv("ENS_APP_MODE", "ambos"))

def main() -> None:
    """
    Executa o aplicativo com preferencia por Waitress no Windows e fallback para o servidor do Flask.
    """
    modo = os.getenv("ENS_APP_MODE", "ambos")
    modo_normalizado = (modo or "ambos").strip().lower()
    if modo_normalizado not in ("admin", "publico", "ambos"):
        print(f"Valor invalido para ENS_APP_MODE={modo!r}; usando 'ambos'.", file=sys.stderr)
        modo_normalizado = "ambos"

    global app
    aplicativo = app
    if aplicativo is None or aplicativo.config.get("ENS_APP_MODE") != modo_normalizado:
        aplicativo = criar_aplicativo(modo_normalizado)
        app = aplicativo

    host = os.getenv("ENS_HOST", "0.0.0.0")
    porta_valor = os.getenv("ENS_PORT", "8080")
    try:
        port = int(porta_valor)
    except (TypeError, ValueError):
        print(f"Valor invalido para ENS_PORT={porta_valor!r}; usando 8080.", file=sys.stderr)
        port = 8080

    display_host = "127.0.0.1" if host in ("0.0.0.0", "::") else host
    use_waitress = (os.name == "nt") or env_flag("ENS_USE_WAITRESS", True)
    url = f"http://{display_host}:{port}"

    threads_env = os.getenv("ENS_WAITRESS_THREADS")
    waitress_threads: int | None = None
    if threads_env:
        try:
            candidate = int(threads_env)
            if candidate > 0:
                waitress_threads = candidate
            else:
                print(f"ENS_WAITRESS_THREADS deve ser maior que zero (valor atual: {threads_env!r}).", file=sys.stderr)
        except ValueError:
            print(f"Valor invalido para ENS_WAITRESS_THREADS={threads_env!r}; ignorando configuracao.", file=sys.stderr)

    if use_waitress:
        try:
            from waitress import serve

            serve_kwargs: dict[str, object] = {"host": host, "port": port}
            if waitress_threads:
                serve_kwargs["threads"] = waitress_threads
            log_extra = f" com {waitress_threads} threads" if waitress_threads else ""
            print(f"[waitress] Servidor ENS ouvindo em {url} (bind {host}:{port}){log_extra}")
            serve(aplicativo, **serve_kwargs)
            return
        except Exception as exc:
            print("Waitress indisponivel, utilizando servidor de desenvolvimento Flask:", exc, file=sys.stderr)

    print(f"[flask] Servidor ENS ouvindo em {url} (bind {host}:{port})")
    aplicativo.run(host=host, port=port, debug=False, use_reloader=False)


if __name__ == "__main__":
    main()
